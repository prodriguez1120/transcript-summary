import os
import json
import re
from openai import OpenAI
from typing import List, Dict, Any, Tuple
from pathlib import Path
import time
from dotenv import load_dotenv
import concurrent.futures
from functools import lru_cache
import hashlib
from datetime import datetime
from prompt_config import get_prompt_config

# Import configuration
try:
    from config import get_output_path, OUTPUT_FILES

    CONFIG_AVAILABLE = True
except ImportError:
    CONFIG_AVAILABLE = False
    print("Warning: config.py not found, using default output directory")

# Check if python-docx is available
try:
    from docx import Document
    from docx.shared import Inches

    DOCX_AVAILABLE = True
    print("python-docx library is available")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"python-docx library not available: {e}")

# Check if ChromaDB is available
try:
    import chromadb
    from chromadb.config import Settings

    CHROMA_AVAILABLE = True
    print("ChromaDB library is available")
except ImportError as e:
    CHROMA_AVAILABLE = False
    print(f"ChromaDB library not available: {e}")

# Load environment variables
load_dotenv()


class TranscriptSummarizer:
    def __init__(
        self,
        api_key: str = None,
        max_workers: int = 3,
        chroma_persist_directory: str = "./chroma_db",
    ):
        """Initialize the transcript summarizer with OpenAI API key and ChromaDB."""
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError(
                "OpenAI API key is required. Set OPENAI_API_KEY environment variable or pass it to the constructor."
            )

        self.client = OpenAI(api_key=self.api_key)
        self.max_workers = max_workers  # Limit concurrent API calls
        self.chroma_persist_directory = chroma_persist_directory

        # Initialize ChromaDB if available
        if CHROMA_AVAILABLE:
            try:
                self.chroma_client = chromadb.PersistentClient(
                    path=chroma_persist_directory
                )
                self.collection = self.chroma_client.get_or_create_collection(
                    name="transcript_chunks", metadata={"hnsw:space": "cosine"}
                )
                print(f"ChromaDB initialized at {chroma_persist_directory}")
            except Exception as e:
                print(f"Warning: ChromaDB initialization failed: {e}")
                self.chroma_client = None
                self.collection = None
        else:
            self.chroma_client = None
            self.collection = None
            print("Warning: ChromaDB not available. Vector storage will be disabled.")

        # Semantic chunking parameters
        self.chunk_size = 2000  # Smaller chunks for better semantic coherence
        self.overlap_size = 200  # Overlap for context preservation
        self.max_chunks_per_transcript = 50  # Limit chunks to prevent overwhelming

        # Define the analysis questions with structured numbering
        self.questions = {
            "company_overview": {
                "title": "Company & Interviewee Overview",
                "questions": [
                    "Section 1a: What is the name of the expert (interviewee)?",
                    "Section 1b: What is the name of the relevant organization that they work(ed) for?",
                    "Section 1c: What is/was their title and role?",
                    "Section 1d: What is the company's size (employees, revenue, number of locations if mentioned)?",
                    "Section 1e: What type of services do they offer? (e.g. foreign material detection, quality control, food safety inspection, contamination prevention). Note approximate splits if available",
                ],
            },
            "flexxray_analysis": {
                "title": "FlexXray Analysis",
                "questions": [
                    "Section 2a: What is FlexXray's value proposition to customers?",
                    "Section 2b: What are FlexXray's key strengths and competitive advantages?",
                    "Section 2c: What are FlexXray's main weaknesses and areas of concern?",
                    "Section 2d: What areas require further due diligence or investigation?",
                ],
            },
            "financial_performance": {
                "title": "Financial Performance & Growth",
                "questions": [
                    "Section 3a: What has driven FlexXray's historical growth?",
                    "Section 3b: What are the key drivers and opportunities for FlexXray's future growth?",
                    "Section 3c: What are the key customer purchasing criteria for FlexXray products/services?",
                ],
            },
            "insourcing_risk": {
                "title": "Insourcing Risk & Mitigation",
                "questions": [
                    "Section 4a: What is the insourcing risk that FlexXray faces from their customers?",
                    "Section 4b: What factors drive customers to consider insourcing?",
                    "Section 4c: How can FlexXray mitigate the risk of customer insourcing?",
                ],
            },
            "porters_five_forces": {
                "title": "Porter's 5 Forces Analysis",
                "questions": [
                    "Section 5a: What is the threat of new entrants to FlexXray's market?",
                    "Section 5b: What is the bargaining power of suppliers in FlexXray's industry?",
                    "Section 5c: What is the bargaining power of customers/buyers?",
                    "Section 5d: What is the threat of substitute products or services?",
                    "Section 5e: What is the intensity of competitive rivalry in the food contamination inspection market?",
                ],
            },
            "strategic_analysis": {
                "title": "Strategic Analysis: Growth, Risks & Industry Trends",
                "questions": [
                    "Section 6a: What growth opportunities or expansion plans are discussed?",
                    "Section 6b: What new markets, products, or services is FlexXray pursuing?",
                    "Section 6c: What strategic initiatives or investments are being made?",
                    "Section 6d: What are the main risks or concerns mentioned about FlexXray's business?",
                    "Section 6e: What key industry trends are affecting the food contamination inspection market?",
                    "Section 6f: What technological changes are disrupting or enabling the industry?",
                ],
            },
        }

    @lru_cache(maxsize=128)
    def extract_text_from_document(self, doc_path: str) -> str:
        """Extract text content from a Word document with caching for repeated calls."""
        try:
            file_extension = Path(doc_path).suffix.lower()

            if file_extension == ".docx":
                # Handle Word documents - optimized for transcript processing
                doc = Document(doc_path)
                text_parts = []

                # Extract text from paragraphs (most efficient for transcripts)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        text_parts.append(paragraph.text.strip())

                # Join with double newlines for better readability
                return "\n\n".join(text_parts)
            else:
                print(
                    f"Unsupported file format: {file_extension}. Only .docx files are supported."
                )
                return ""

        except Exception as e:
            print(f"Error extracting text from {doc_path}: {e}")
            return ""

    def semantic_chunk_text(
        self, text: str, transcript_name: str, doc_path: str = None
    ) -> List[Dict[str, Any]]:
        """
        Create intelligent semantic chunks with logical break awareness for expert call transcripts.

        This enhanced chunking strategy:
        1. Identifies logical breaks (paragraphs, topic transitions, speaker changes)
        2. Uses dynamic chunk sizing based on content structure
        3. Preserves semantic coherence by avoiding cutting mid-topic
        4. Improves overlap handling for better context preservation
        """
        if not text or len(text.strip()) == 0:
            return []

        # Clean and normalize text
        text = re.sub(r"\s+", " ", text).strip()

        # Enhanced text segmentation with logical break detection
        segments = self._segment_text_with_logical_breaks(text)

        chunks = []
        chunk_id = 0

        for segment in segments:
            if len(segment.strip()) == 0:
                continue

            # Check if segment is too long and needs further splitting
            if len(segment) > self.chunk_size:
                # Split long segments at sentence boundaries
                sub_chunks = self._split_long_segment(
                    segment, transcript_name, chunk_id, doc_path
                )
                chunks.extend(sub_chunks)
                chunk_id += len(sub_chunks)
            else:
                # Create chunk from appropriately sized segment
                chunk_data = self._create_chunk_metadata(
                    segment.strip(), transcript_name, chunk_id, doc_path, chunk_id
                )
                chunks.append(chunk_data)
                chunk_id += 1

            # Check if we've reached the maximum chunks limit
            if chunk_id >= self.max_chunks_per_transcript:
                break

        # Add overlap between consecutive chunks for better context
        chunks = self._add_intelligent_overlap(chunks)

        print(
            f"Created {len(chunks)} intelligent semantic chunks for {transcript_name}"
        )
        return chunks

    def _segment_text_with_logical_breaks(self, text: str) -> List[str]:
        """
        Segment text at logical breaks including:
        - Paragraph breaks (double newlines)
        - Speaker changes (Interviewer/Expert patterns)
        - Topic transitions (key phrases indicating new subjects)
        - Natural section breaks
        """
        segments = []

        # Split by paragraph breaks first
        paragraphs = re.split(r"\n\s*\n", text)

        for paragraph in paragraphs:
            if len(paragraph.strip()) == 0:
                continue

            # Check for speaker changes within paragraphs
            speaker_segments = self._split_by_speaker_changes(paragraph)
            segments.extend(speaker_segments)

        # Further segment by topic transitions
        final_segments = []
        for segment in segments:
            if len(segment.strip()) == 0:
                continue

            # Check if segment contains topic transitions
            topic_segments = self._split_by_topic_transitions(segment)
            final_segments.extend(topic_segments)

        return final_segments

    def _split_by_speaker_changes(self, text: str) -> List[str]:
        """Split text by speaker changes (Interviewer/Expert patterns)."""
        segments = []

        # Common speaker patterns in expert calls
        speaker_patterns = [
            r"(Interviewer:|Expert:|Q:|A:)",
            r"(Interviewer\s*:)",
            r"(Expert\s*:)",
            r"(Q\s*:)",
            r"(A\s*:)",
            r"([A-Z][a-z]+\s*:)",  # Names followed by colon
        ]

        # Find all speaker change points
        change_points = []
        for pattern in speaker_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                change_points.append(match.start())

        # Sort change points and add start/end
        change_points = sorted(list(set(change_points)))
        change_points = [0] + change_points + [len(text)]

        # Create segments between change points
        for i in range(len(change_points) - 1):
            start = change_points[i]
            end = change_points[i + 1]
            segment = text[start:end].strip()

            if len(segment) > 0:
                segments.append(segment)

        return segments if segments else [text]

    def _split_by_topic_transitions(self, text: str) -> List[str]:
        """Split text by topic transition indicators."""
        segments = []

        # Topic transition phrases that indicate new subjects
        topic_transitions = [
            r"\b(Now|Next|Moving on|Let\'s talk about|Regarding|As for|In terms of|When it comes to)\b",
            r"\b(First|Second|Third|Finally|Additionally|Moreover|Furthermore|However|On the other hand)\b",
            r"\b(Let me ask|Can you tell me|What about|How about|Tell me about)\b",
            r"\b(So|Therefore|Thus|Hence|As a result|Consequently)\b",
            r"\b(For example|For instance|Specifically|In particular|To illustrate)\b",
            r"\b(In conclusion|To summarize|Overall|In summary)\b",
        ]

        # Find topic transition points
        transition_points = []
        for pattern in topic_transitions:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Only consider transitions that aren't at the very beginning
                if match.start() > 50:  # Minimum distance from start
                    transition_points.append(match.start())

        # Sort transition points
        transition_points = sorted(list(set(transition_points)))

        if not transition_points:
            return [text]

        # Create segments between transitions
        start = 0
        for point in transition_points:
            segment = text[start:point].strip()
            if len(segment) > 0:
                segments.append(segment)
            start = point

        # Add final segment
        final_segment = text[start:].strip()
        if len(final_segment) > 0:
            segments.append(final_segment)

        return segments if segments else [text]

    def _split_long_segment(
        self, segment: str, transcript_name: str, base_chunk_id: int, doc_path: str
    ) -> List[Dict[str, Any]]:
        """Split long segments at sentence boundaries while preserving context."""
        chunks = []

        # Split into sentences
        sentences = re.split(r"(?<=[.!?])\s+", segment)

        current_chunk = ""
        chunk_id = base_chunk_id

        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue

            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Create chunk
                chunk_data = self._create_chunk_metadata(
                    current_chunk.strip(), transcript_name, chunk_id, doc_path, chunk_id
                )
                chunks.append(chunk_data)

                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(
                    current_chunk, self.overlap_size
                )
                current_chunk = overlap_sentences + " " + sentence
                chunk_id += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        # Add the last chunk if there's remaining content
        if current_chunk.strip():
            chunk_data = self._create_chunk_metadata(
                current_chunk.strip(), transcript_name, chunk_id, doc_path, chunk_id
            )
            chunks.append(chunk_data)

        return chunks

    def _add_intelligent_overlap(
        self, chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Add intelligent overlap between chunks for better context preservation."""
        if len(chunks) <= 1:
            return chunks

        enhanced_chunks = []

        for i, chunk in enumerate(chunks):
            if i == 0:
                # First chunk: no overlap needed
                enhanced_chunks.append(chunk)
                continue

            # Get overlap from previous chunk
            previous_text = chunks[i - 1]["text"]
            overlap_text = self._get_overlap_sentences(previous_text, self.overlap_size)

            # Add overlap to current chunk if it doesn't already contain it
            current_text = chunk["text"]
            if overlap_text and overlap_text not in current_text:
                # Find the best place to insert overlap (usually at the beginning)
                enhanced_text = overlap_text + " " + current_text

                # Update chunk with enhanced text
                enhanced_chunk = chunk.copy()
                enhanced_chunk["text"] = enhanced_text
                enhanced_chunk["metadata"] = chunk["metadata"].copy()
                enhanced_chunk["metadata"]["chunk_length"] = len(enhanced_text)
                enhanced_chunk["metadata"]["has_overlap"] = True

                enhanced_chunks.append(enhanced_chunk)
            else:
                enhanced_chunks.append(chunk)

        return enhanced_chunks

    def _create_chunk_metadata(
        self,
        text: str,
        transcript_name: str,
        chunk_id: int,
        doc_path: str,
        sentence_position: int,
    ) -> Dict[str, Any]:
        """Create metadata for a chunk including semantic context."""
        # Generate unique ID for the chunk
        chunk_hash = hashlib.md5(
            f"{transcript_name}_{chunk_id}_{text[:100]}".encode()
        ).hexdigest()

        # Extract key entities and topics for better searchability
        key_topics = self._extract_key_topics(text)

        # Convert key_topics list to string for ChromaDB compatibility
        key_topics_str = ", ".join(key_topics) if key_topics else ""

        # Create metadata
        metadata = {
            "transcript_name": transcript_name,
            "chunk_id": chunk_id,
            "chunk_hash": chunk_hash,
            "doc_path": doc_path,
            "sentence_position": sentence_position,
            "chunk_length": len(text),
            "key_topics": key_topics_str,
            "timestamp": datetime.now().isoformat(),
            "chunk_type": self._classify_chunk_type(text),
        }

        return {"id": chunk_hash, "text": text, "metadata": metadata}

    def _extract_key_topics(self, text: str) -> List[str]:
        """Extract key topics and entities from chunk text."""
        topics = []
        text_lower = text.lower()

        # Look for company names, products, and key terms
        key_terms = [
            "flexxray",
            "x-ray",
            "xray",
            "foreign material",
            "contamination",
            "food safety",
            "quality control",
            "inspection",
            "detection",
            "customer",
            "market",
            "industry",
            "growth",
            "revenue",
            "profit",
            "competition",
            "competitive",
            "advantage",
            "strength",
            "weakness",
            "risk",
            "opportunity",
            "technology",
            "innovation",
            "investment",
        ]

        for term in key_terms:
            if term in text_lower:
                topics.append(term)

        # Look for company names (capitalized words that might be company names)
        company_patterns = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", text)
        for company in company_patterns[:3]:  # Limit to first 3 potential companies
            if len(company.split()) <= 3 and company.lower() not in [
                "the",
                "and",
                "for",
                "with",
            ]:
                topics.append(company)

        return list(set(topics))  # Remove duplicates

    def _classify_chunk_type(self, text: str) -> str:
        """
        Intelligently classify the type of content in the chunk based on semantic analysis.
        Enhanced for expert call transcripts with better pattern recognition.
        """
        text_lower = text.lower()

        # Financial and business metrics
        financial_indicators = [
            "financial",
            "revenue",
            "profit",
            "cost",
            "pricing",
            "budget",
            "margin",
            "earnings",
            "cash flow",
            "investment",
            "funding",
            "valuation",
            "market cap",
            "ebitda",
            "quarterly",
            "annual",
            "growth rate",
            "market share",
            "pricing strategy",
        ]

        # Customer and market analysis
        customer_market_indicators = [
            "customer",
            "client",
            "market",
            "demand",
            "sales",
            "customer satisfaction",
            "market demand",
            "customer needs",
            "market trends",
            "customer relationships",
            "market size",
            "customer acquisition",
            "market penetration",
            "customer retention",
        ]

        # Technology and product development
        technology_indicators = [
            "technology",
            "innovation",
            "product",
            "development",
            "research",
            "r&d",
            "technical",
            "engineering",
            "software",
            "hardware",
            "platform",
            "solution",
            "product features",
            "technical capabilities",
            "product roadmap",
            "innovation",
        ]

        # Competitive analysis and positioning
        competitive_indicators = [
            "competition",
            "competitive",
            "advantage",
            "strength",
            "weakness",
            "differentiation",
            "market position",
            "competitive landscape",
            "competitive advantage",
            "market leader",
            "competitive analysis",
            "competitive pressure",
            "market positioning",
        ]

        # Risk and challenges
        risk_indicators = [
            "risk",
            "challenge",
            "problem",
            "issue",
            "concern",
            "threat",
            "obstacle",
            "regulatory",
            "compliance",
            "legal",
            "operational risk",
            "market risk",
            "competitive risk",
            "technology risk",
            "financial risk",
        ]

        # Growth and strategy
        growth_indicators = [
            "growth",
            "expansion",
            "opportunity",
            "future",
            "strategy",
            "strategic",
            "market expansion",
            "product development",
            "strategic initiatives",
            "growth strategy",
            "expansion plans",
            "market opportunity",
            "strategic positioning",
        ]

        # Industry and market trends
        industry_indicators = [
            "industry",
            "sector",
            "market trends",
            "industry trends",
            "regulatory changes",
            "market dynamics",
            "industry dynamics",
            "market evolution",
            "industry evolution",
        ]

        # Operations and processes
        operational_indicators = [
            "operations",
            "process",
            "efficiency",
            "quality",
            "supply chain",
            "logistics",
            "manufacturing",
            "production",
            "operational excellence",
            "process improvement",
        ]

        # Calculate scores for each category
        scores = {
            "financial": sum(
                1 for indicator in financial_indicators if indicator in text_lower
            ),
            "customer_market": sum(
                1 for indicator in customer_market_indicators if indicator in text_lower
            ),
            "technology_product": sum(
                1 for indicator in technology_indicators if indicator in text_lower
            ),
            "competitive_analysis": sum(
                1 for indicator in competitive_indicators if indicator in text_lower
            ),
            "risk_challenge": sum(
                1 for indicator in risk_indicators if indicator in text_lower
            ),
            "growth_strategy": sum(
                1 for indicator in growth_indicators if indicator in text_lower
            ),
            "industry_trends": sum(
                1 for indicator in industry_indicators if indicator in text_lower
            ),
            "operational": sum(
                1 for indicator in operational_indicators if indicator in text_lower
            ),
        }

        # Find the category with the highest score
        best_category = max(scores, key=scores.get)

        # Only return a specific category if there's a meaningful score
        if scores[best_category] >= 2:
            return best_category
        elif scores[best_category] >= 1:
            # Check if it's a mixed category
            mixed_categories = [cat for cat, score in scores.items() if score >= 1]
            if len(mixed_categories) > 1:
                return "mixed_" + "_".join(mixed_categories[:2])
            else:
                return best_category
        else:
            return "general"

    def _get_overlap_sentences(self, text: str, overlap_size: int) -> str:
        """Get the last few sentences for overlap context."""
        sentences = text.split(".")
        overlap_text = ""

        # Start from the end and work backwards
        for sentence in reversed(sentences):
            if len(overlap_text + sentence) <= overlap_size:
                overlap_text = sentence + ". " + overlap_text
            else:
                break

        return overlap_text.strip()

    def store_chunks_in_vector_db(self, chunks: List[Dict[str, Any]]) -> bool:
        """Store chunks in ChromaDB vector database."""
        if not self.collection or not CHROMA_AVAILABLE:
            print("Warning: ChromaDB not available. Skipping vector storage.")
            return False

        try:
            # Prepare data for ChromaDB
            ids = [chunk["id"] for chunk in chunks]
            texts = [chunk["text"] for chunk in chunks]
            metadatas = [chunk["metadata"] for chunk in chunks]

            # Add to collection
            self.collection.add(ids=ids, documents=texts, metadatas=metadatas)

            print(f"Stored {len(chunks)} chunks in ChromaDB")
            return True

        except Exception as e:
            print(f"Error storing chunks in ChromaDB: {e}")
            return False

    def search_relevant_chunks(
        self, query: str, transcript_name: str = None, top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """Search for relevant chunks based on semantic similarity."""
        if not self.collection or not CHROMA_AVAILABLE:
            print("Warning: ChromaDB not available. Cannot search for relevant chunks.")
            return []

        try:
            # Build search query
            search_query = query

            # Add transcript filter if specified
            where_filter = {}
            if transcript_name:
                where_filter["transcript_name"] = transcript_name

            # Search in ChromaDB
            results = self.collection.query(
                query_texts=[search_query],
                n_results=top_k,
                where=where_filter if where_filter else None,
            )

            # Format results
            relevant_chunks = []
            if results["documents"] and results["documents"][0]:
                for i, doc in enumerate(results["documents"][0]):
                    chunk_data = {
                        "text": doc,
                        "metadata": (
                            results["metadatas"][0][i]
                            if results["metadatas"] and results["metadatas"][0]
                            else {}
                        ),
                        "distance": (
                            results["distances"][0][i]
                            if results["distances"] and results["distances"][0]
                            else None
                        ),
                    }
                    relevant_chunks.append(chunk_data)

            return relevant_chunks

        except Exception as e:
            print(f"Error searching ChromaDB: {e}")
            return []

    def get_context_for_question(
        self,
        query: str,
        transcript_name: str = None,
        top_k: int = 3,
        use_hierarchical: bool = False,
        top_k_per_transcript: int = 5,
    ) -> str:
        """Get relevant context chunks for a specific question."""
        if not self.collection or not CHROMA_AVAILABLE:
            print("Warning: ChromaDB not available. Cannot get context for question.")
            return ""

        try:
            if use_hierarchical:
                # Use enhanced hierarchical retrieval
                relevant_chunks = self.hierarchical_retrieval_with_reranking(
                    query, top_k_per_transcript=top_k_per_transcript, final_top_k=top_k
                )
            else:
                # Use original simple search
                relevant_chunks = self.search_relevant_chunks(
                    query, transcript_name, top_k
                )

            if not relevant_chunks:
                return ""

            # Combine relevant chunks with context
            context_parts = []
            for chunk in relevant_chunks:
                if isinstance(chunk, dict) and "text" in chunk:
                    if "rank" in chunk:
                        # Hierarchical retrieval result
                        context_parts.append(
                            f"Rank {chunk['rank']} (Score: {chunk.get('final_score', 'N/A'):.3f}) "
                            f"from {chunk.get('transcript_name', 'Unknown')} "
                            f"[{chunk.get('chunk_type', 'unknown')}]:\n{chunk['text']}"
                        )
                    else:
                        # Simple search result
                        context_parts.append(
                            f"Context from {chunk['metadata'].get('transcript_name', 'Unknown')}:\n{chunk['text']}"
                        )

            return "\n\n".join(context_parts)

        except Exception as e:
            print(f"Error getting context for question: {e}")
            return ""

    def analyze_section(
        self,
        transcript_text: str,
        section_key: str,
        section_data: dict,
        transcript_name: str,
        use_hierarchical_context: bool = False,
    ) -> tuple:
        """Analyze a single section of a transcript with semantic context from vector store."""
        try:
            # Get relevant context from vector store for this question
            section_questions = "\n".join(
                [f"{i+1}. {q}" for i, q in enumerate(section_data["questions"])]
            )

            # Search for relevant context chunks (optionally using hierarchical retrieval)
            relevant_context = self.get_context_for_question(
                section_data["title"] + " " + section_questions,
                transcript_name,
                top_k=3,
                use_hierarchical=use_hierarchical_context,
                top_k_per_transcript=5,
            )

            # Create the main prompt with context
            context_section = ""
            if relevant_context:
                context_section = f"""
RELEVANT CONTEXT FROM PREVIOUS ANALYSES:
{relevant_context}

"""

            # Get prompt configuration
            prompt_config = get_prompt_config()

            main_prompt = prompt_config.format_prompt(
                "transcript_analysis",
                context_section=context_section,
                transcript_text=transcript_text,
                section_title=section_data["title"],
                section_questions=section_questions,
            )

            full_prompt = (
                f"{main_prompt}\n\n{section_data['title']}:\n{section_questions}"
            )

            # Get OpenAI parameters from config
            params = prompt_config.get_prompt_parameters("transcript_analysis")

            response = self.client.chat.completions.create(
                model=params.get("model", "gpt-4o"),
                messages=[
                    {
                        "role": "system",
                        "content": prompt_config.get_system_message(
                            "transcript_analysis"
                        ),
                    },
                    {"role": "user", "content": full_prompt},
                ],
                temperature=params.get("temperature", 0.3),
                top_p=params.get("top_p", 1.0),
                max_tokens=params.get("max_tokens", 4000),
            )

            # Get the main analysis
            main_analysis = response.choices[0].message.content.strip()

            # Extract and select best quotes across all experts for this section
            best_quotes = self.extract_and_select_best_quotes(
                section_data["title"] + " " + section_questions,
                self.results if hasattr(self, "results") else [],
                max_quotes_per_expert=4,
                max_final_quotes=3,
            )

            # Add selected quotes to the response
            if best_quotes:
                quotes_section = "\n\nSELECTED EXPERT QUOTES:\n" + "=" * 50 + "\n"
                for i, quote_data in enumerate(best_quotes, 1):
                    quotes_section += f"\n{i}. \"{quote_data['quote']}\""
                    quotes_section += f"\n   -- {quote_data['expert']}"
                    quotes_section += f"\n   (Relevance: {quote_data['relevance_score']:.3f}, Type: {quote_data['chunk_type']})\n"

                full_response = main_analysis + quotes_section
            else:
                full_response = main_analysis

            return (
                section_key,
                {"title": section_data["title"], "answers": full_response},
            )

        except Exception as e:
            print(f"Error analyzing {section_key} for {transcript_name}: {e}")
            return (
                section_key,
                {
                    "title": section_data["title"],
                    "answers": f"Error during analysis: {e}",
                },
            )

    def analyze_transcript(
        self, transcript_text: str, transcript_name: str
    ) -> Dict[str, Any]:
        """Analyze a single transcript using parallel processing for sections."""
        results = {
            "transcript_name": transcript_name,
            "transcript_text": transcript_text,
            "analysis": {},
        }

        # Process sections in parallel with limited concurrency
        with concurrent.futures.ThreadPoolExecutor(
            max_workers=self.max_workers
        ) as executor:
            # Submit all section analysis tasks
            future_to_section = {
                executor.submit(
                    self.analyze_section,
                    transcript_text,
                    section_key,
                    section_data,
                    transcript_name,
                ): section_key
                for section_key, section_data in self.questions.items()
            }

            # Collect results as they complete
            for future in concurrent.futures.as_completed(future_to_section):
                section_key, section_result = future.result()
                results["analysis"][section_key] = section_result

                # Small delay to avoid rate limiting
                time.sleep(0.5)

        return results

    def process_transcripts_directory(
        self, directory_path: str, progress_callback=None
    ) -> List[Dict[str, Any]]:
        """Process all Word document transcripts in a directory with semantic chunking and vector storage."""
        directory = Path(directory_path)
        docx_files = list(directory.glob("*.docx"))

        if not docx_files:
            print(f"No Word documents (.docx) found in {directory_path}")
            return []

        all_results = []
        total_files = len(docx_files)

        # Process transcripts sequentially but sections in parallel
        for i, doc_file in enumerate(docx_files):
            if progress_callback:
                progress = (i / total_files) * 100
                progress_callback(
                    f"Processing {doc_file.name} ({i+1}/{total_files})", progress
                )

            print(f"\nProcessing: {doc_file.name}")

            # Extract text from document
            transcript_text = self.extract_text_from_document(str(doc_file))

            if not transcript_text:
                print(f"Could not extract text from {doc_file.name}")
                continue

            # Create semantic chunks and store in vector database
            print(f"Creating semantic chunks for {doc_file.name}...")
            chunks = self.semantic_chunk_text(
                transcript_text, doc_file.name, str(doc_file)
            )
            print(f"Created {len(chunks)} semantic chunks")

            # Store chunks in vector database
            if chunks and self.collection:
                print(f"Storing chunks in vector database...")
                self.store_chunks_in_vector_db(chunks)

            # Analyze the transcript with parallel section processing
            results = self.analyze_transcript(transcript_text, doc_file.name)
            # Add document path for enhanced citation extraction
            results["doc_path"] = str(doc_file)
            # Add chunk information
            results["chunks"] = chunks
            all_results.append(results)

            print(f"Completed analysis of {doc_file.name}")

        # Store results in class for quote extraction
        self.results = all_results

        if progress_callback:
            progress_callback("Analysis complete!", 100)

        return all_results

    def save_results(self, results: List[Dict[str, Any]], output_file: str = None):
        """Save analysis results to a JSON file."""
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(OUTPUT_FILES["transcript_analysis_json"])
            else:
                output_file = "FlexXray_transcript_analysis_results.json"

        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"Results saved to {output_file}")

    def _extract_citation_enhanced(self, filename: str, doc_path: str = None) -> str:
        """Enhanced citation extraction that tries filename first, then document content."""
        # First try to extract from filename
        citation = self._extract_citation_from_filename(filename)

        # If we got a generic result (like just the filename), try document content
        if citation == filename or len(citation.split()) < 2:
            if doc_path:
                doc_citation = self._extract_expert_name_from_document(doc_path)
                if doc_citation:
                    return doc_citation

        return citation

    def _extract_citation_from_filename(self, filename: str) -> str:
        """Extract a clean expert name from the filename."""
        # Remove file extension
        name_without_ext = filename.replace(".docx", "").replace(".pdf", "")

        # Try to extract expert name from common patterns

        # Pattern 1: First_Last- (underscore and dash)
        match = re.search(r"^([A-Za-z]+_[A-Za-z]+)", name_without_ext)
        if match:
            return match.group(1).replace("_", " ")

        # Pattern 2: First Last - (space and dash)
        match = re.search(r"^([A-Za-z]+\s+[A-Za-z]+)", name_without_ext)
        if match:
            return match.group(1)

        # Pattern 3: "Third Bridge Call - Name - ..."
        match = re.search(r"Third Bridge Call - ([^-]+)", name_without_ext)
        if match:
            return match.group(1).strip()

        # Pattern 4: "Name - Current ..."
        match = re.search(r"^([^-]+) - Current", name_without_ext)
        if match:
            return match.group(1).strip()

        # Pattern 5: "Name - Company - ..."
        match = re.search(r"^([^-]+) - [^-]+ -", name_without_ext)
        if match:
            return match.group(1).strip()

        # Pattern 6: "Name, Title - ..."
        match = re.search(r"^([^,]+),", name_without_ext)
        if match:
            return match.group(1).strip()

        # Fallback: use first two words of the filename
        words = name_without_ext.split()
        if len(words) >= 2:
            return f"{words[0]} {words[1]}"
        elif len(words) == 1:
            return words[0]
        else:
            return filename

    def _extract_expert_name_from_document(self, doc_path: str) -> str:
        """Extract expert name from document content when not found in filename."""
        try:
            # Extract text from document
            doc_text = self.extract_text_from_document(doc_path)

            # Look for expert name patterns in the first 1000 characters
            first_part = doc_text[:1000].lower()

            # Pattern 1: "Interview with [Name]"
            match = re.search(r"interview with ([a-z]+\s+[a-z]+)", first_part)
            if match:
                return match.group(1).title()

            # Pattern 2: "Expert: [Name]"
            match = re.search(r"expert:\s*([a-z]+\s+[a-z]+)", first_part)
            if match:
                return match.group(1).title()

            # Pattern 3: "Interviewee: [Name]"
            match = re.search(r"interviewee:\s*([a-z]+\s+[a-z]+)", first_part)
            if match:
                return match.group(1).title()

            # Pattern 4: "Name - Title" at the beginning
            match = re.search(r"^([A-Z][a-z]+\s+[A-Z][a-z]+)\s*[-–]\s*", doc_text[:500])
            if match:
                return match.group(1)

            # Pattern 5: Look for capitalized name patterns in first few lines
            lines = doc_text.split("\n")[:10]
            for line in lines:
                # Look for "First Last" pattern with proper capitalization
                match = re.search(r"\b([A-Z][a-z]+\s+[A-Z][a-z]+)\b", line)
                if match:
                    name = match.group(1)
                    # Skip common words that might be mistaken for names
                    if name.lower() not in [
                        "the company",
                        "the industry",
                        "the market",
                        "the business",
                    ]:
                        return name

            return None

        except Exception as e:
            print(f"Error extracting expert name from document: {e}")
            return None

    def _extract_quotes_from_answer(self, answer_text: str) -> List[str]:
        """Extract quotes from AI-generated answer text."""
        try:
            quotes = []

            # Look for quotes in the format "Quote: [quote text]"
            quote_matches = re.findall(r'Quote:\s*"([^"]+)"', answer_text)
            quotes.extend(quote_matches)

            # Also look for quotes in the format "Quote: [quote text]" without quotes
            quote_matches_no_quotes = re.findall(r'Quote:\s*([^"\n]+)', answer_text)
            quotes.extend(quote_matches_no_quotes)

            # Remove duplicates and clean up
            unique_quotes = []
            for quote in quotes:
                quote = quote.strip()
                if (
                    quote and quote not in unique_quotes and len(quote) > 10
                ):  # Minimum length to be meaningful
                    unique_quotes.append(quote)

            return unique_quotes[:2]  # Limit to 2 quotes max

        except Exception as e:
            print(f"Error extracting quotes from answer: {e}")
            return []

    def _is_information_unavailable(self, answer_text: str) -> bool:
        """Check if the answer indicates information is not available."""
        unavailable_phrases = [
            "information not available",
            "information is not available",
            "no information available",
            "not available in transcript",
            "not mentioned in transcript",
            "not discussed in transcript",
            "not provided in transcript",
            "not available",
            "not mentioned",
            "not discussed",
            "not provided",
            "n/a",
            "none",
            "unknown",
        ]

        answer_lower = answer_text.lower().strip()

        # Check if the answer is just one of these phrases
        if answer_lower in unavailable_phrases:
            return True

        # Check if the answer starts with these phrases
        for phrase in unavailable_phrases:
            if answer_lower.startswith(phrase):
                return True

        # Check if the answer is very short and contains these phrases
        if (
            len(answer_lower) < 50
        ):  # Short answers are more likely to be "not available"
            for phrase in unavailable_phrases:
                if phrase in answer_lower:
                    return True

        return False

    def _parse_section_answers(
        self, answers: str, questions: List[str]
    ) -> Dict[str, str]:
        """Parse section answers into individual question responses."""

        parsed_answers = {}

        # First, let's try to split the entire response into sections based on common patterns
        # This handles cases where the AI returns all questions in one response

        # Strategy 1: Split by numbered sections with markdown headers
        sections = []

        # Try to find all section headers and split accordingly
        section_headers = re.findall(
            r"(?:####|###)\s*(\d+)\.\s*Section\s+(\d+[a-f]):", answers
        )
        if section_headers:
            # Split by section headers
            parts = re.split(r"(?:####|###)\s*\d+\.\s*Section\s+\d+[a-f]:", answers)
            if len(parts) > 1:
                # Skip the first part (content before first header)
                for i, (match_num, section_num) in enumerate(section_headers):
                    if i + 1 < len(parts):
                        content = parts[i + 1].strip()
                        if content:
                            sections.append((section_num, content))

        # Strategy 2: If no markdown headers, try simple numbered sections
        if not sections:
            numbered_sections = re.findall(r"(\d+)\.\s*Section\s+(\d+[a-f]):", answers)
            if numbered_sections:
                parts = re.split(r"\d+\.\s*Section\s+\d+[a-f]:", answers)
                if len(parts) > 1:
                    for i, (match_num, section_num) in enumerate(numbered_sections):
                        if i + 1 < len(parts):
                            content = parts[i + 1].strip()
                            if content:
                                sections.append((section_num, content))

        # Strategy 3: Look for any "Section Xa:" pattern
        if not sections:
            section_matches = re.findall(r"Section\s+(\d+[a-f]):", answers)
            if section_matches:
                parts = re.split(r"Section\s+\d+[a-f]:", answers)
                if len(parts) > 1:
                    for i, section_num in enumerate(section_matches):
                        if i + 1 < len(parts):
                            content = parts[i + 1].strip()
                            if content:
                                # Find the next section boundary
                                next_section_match = re.search(
                                    r"Section\s+\d+[a-f]:", content
                                )
                                if next_section_match:
                                    content = content[
                                        : next_section_match.start()
                                    ].strip()
                                sections.append((section_num, content))

        # Strategy 4: If still no sections found, try to split by bullet points or numbered lists
        if not sections and len(questions) > 1:
            # Look for bullet points or numbered items that might correspond to questions
            bullet_items = re.findall(
                r"[-•]\s*([^\n]+(?:\n(?![-•])[^\n]*)*)",
                answers,
                re.MULTILINE | re.DOTALL,
            )
            numbered_items = re.findall(
                r"(\d+)\.\s*([^\n]+(?:\n(?!\d+\.)[^\n]*)*)",
                answers,
                re.MULTILINE | re.DOTALL,
            )

            if bullet_items and len(bullet_items) >= len(questions):
                for i, content in enumerate(bullet_items):
                    if i < len(questions):
                        # Try to extract section number from question
                        section_match = re.search(
                            r"Section\s+(\d+[a-f]):", questions[i]
                        )
                        if section_match:
                            sections.append((section_match.group(1), content))

            elif numbered_items and len(numbered_items) >= len(questions):
                for i, (num, content) in enumerate(numbered_items):
                    if i < len(questions):
                        # Try to extract section number from question
                        section_match = re.search(
                            r"Section\s+(\d+[a-f]):", questions[i]
                        )
                        if section_match:
                            sections.append((section_match.group(1), content))

        # Strategy 5: Manual parsing by looking for each question's section number
        if not sections:
            for question in questions:
                section_match = re.search(r"Section\s+(\d+[a-f]):", question)
                if section_match:
                    section_num = section_match.group(1)

                    # Look for this specific section in the answer text
                    patterns = [
                        rf"Section\s+{section_num}:[^\n]*\n(.*?)(?=Section\s+\d+[a-f]:|$)",
                        rf"(?:####|###)\s*\d+\.\s*Section\s+{section_num}:[^\n]*\n(.*?)(?=(?:####|###)\s*\d+\.\s*Section|$)",
                        rf"\d+\.\s*Section\s+{section_num}:[^\n]*\n(.*?)(?=\d+\.\s*Section|$)",
                    ]

                    for pattern in patterns:
                        match = re.search(pattern, answers, re.MULTILINE | re.DOTALL)
                        if match:
                            content = match.group(1).strip()
                            if content:
                                sections.append((section_num, content))
                            break

        # Now process the found sections
        for section_num, content in sections:
            # Find the corresponding question
            target_question = None
            for question in questions:
                if f"Section {section_num}:" in question:
                    target_question = question
                    break

            if target_question:
                answer_text = self._clean_answer_text(content)
                if answer_text and not self._is_information_unavailable(answer_text):
                    parsed_answers[target_question] = answer_text

        # Final fallback: if no structured parsing worked, treat as single response to first question
        if not parsed_answers and questions and answers.strip():
            question = questions[0]
            answer_text = self._clean_answer_text(answers)
            if not self._is_information_unavailable(answer_text):
                parsed_answers[question] = answer_text

        return parsed_answers

    def _clean_answer_text(self, text: str) -> str:
        """Clean and format answer text."""
        if not text:
            return ""

        # Remove bullet points and dashes at the start
        text = re.sub(r"^\s*[-•]\s*", "", text)
        text = re.sub(r"\n\s*[-•]\s*", "\n", text)

        # Remove question numbers and markdown formatting
        text = re.sub(r"^\d+\.\s*\*\*[^*]+\*\*", "", text)
        text = re.sub(r"^\d+\.\s*", "", text)

        # Remove section headers that might be in the answer
        text = re.sub(r"Section\s+\d+[a-f]:[^\n]*\n?", "", text)

        # Remove question text that might be included in the response
        # Look for common question patterns and remove them
        question_patterns = [
            r"What are FlexXray\'s key strengths and competitive advantages\?",
            r"What are FlexXray\'s main weaknesses and areas of concern\?",
            r"What areas require further due diligence or investigation\?",
            r"What has driven FlexXray\'s historical growth\?",
            r"What are the key drivers and opportunities for FlexXray\'s future growth\?",
            r"What are the key customer purchasing criteria for FlexXray products/services\?",
            r"What is the insourcing risk that FlexXray faces from their customers\?",
            r"What factors drive customers to consider insourcing\?",
            r"How can FlexXray mitigate the risk of customer insourcing\?",
            r"What is the threat of new entrants to FlexXray\'s market\?",
            r"What is the bargaining power of suppliers in FlexXray\'s industry\?",
            r"What is the bargaining power of customers/buyers\?",
            r"What is the threat of substitute products or services\?",
            r"What is the intensity of competitive rivalry in the food contamination inspection market\?",
            r"What growth opportunities or expansion plans are discussed\?",
            r"What new markets, products, or services is FlexXray pursuing\?",
            r"What strategic initiatives or investments are being made\?",
            r"What are the main risks or concerns mentioned about FlexXray\'s business\?",
            r"What key industry trends are affecting the food contamination inspection market\?",
            r"What technological changes are disrupting or enabling the industry\?",
        ]

        for pattern in question_patterns:
            # Remove the question text if it appears at the beginning of the response
            text = re.sub(rf"^\s*{pattern}\s*", "", text, flags=re.IGNORECASE)
            # Also remove it if it appears after a bullet point or dash
            text = re.sub(rf"^\s*[-•]\s*{pattern}\s*", "", text, flags=re.IGNORECASE)
            # Remove it if it appears anywhere in the text (but be more careful)
            text = re.sub(rf"\n\s*{pattern}\s*\n", "\n", text, flags=re.IGNORECASE)

        # Remove any remaining question-like patterns
        text = re.sub(r"^\s*[•-]\s*[^:]*\?", "", text)
        text = re.sub(r"\n\s*[•-]\s*[^:]*\?", "\n", text)

        # Clean up extra whitespace and empty lines
        text = re.sub(r"\n\s*\n", "\n", text)
        text = re.sub(r"^\s*\n", "", text)
        text = text.strip()

        return text

    def _consolidate_results(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Consolidate results into the simplified structure."""
        consolidated = {"company_overview_by_expert": {}, "questions_by_section": {}}

        # Process company overview by expert
        for result in results:
            transcript_name = result.get("transcript_name", "Unknown")
            doc_path = result.get("doc_path", None)
            citation = self._extract_citation_enhanced(transcript_name, doc_path)

            if "company_overview" in result.get("analysis", {}):
                overview_data = result["analysis"]["company_overview"]
                consolidated["company_overview_by_expert"][citation] = {
                    "transcript_name": transcript_name,
                    "citation": citation,
                    "answers": overview_data.get("answers", ""),
                    "doc_path": doc_path,
                }

        # Process all other sections by question
        for section_key, section_data in self.questions.items():
            if section_key == "company_overview":
                continue  # Skip company overview as it's handled separately

            section_title = section_data["title"]
            questions = section_data["questions"]

            if section_title not in consolidated["questions_by_section"]:
                consolidated["questions_by_section"][section_title] = {}

            # Process each question in the section
            for question in questions:
                if question not in consolidated["questions_by_section"][section_title]:
                    consolidated["questions_by_section"][section_title][question] = []

                # Collect responses from all transcripts for this question
                for result in results:
                    transcript_name = result.get("transcript_name", "Unknown")
                    doc_path = result.get("doc_path", None)
                    citation = self._extract_citation_enhanced(
                        transcript_name, doc_path
                    )

                    if section_key in result.get("analysis", {}):
                        section_answers = result["analysis"][section_key].get(
                            "answers", ""
                        )
                        parsed_answers = self._parse_section_answers(
                            section_answers, questions
                        )

                        if question in parsed_answers:
                            answer_text = parsed_answers[question]

                            # Extract quotes from the answer text
                            quotes = self._extract_quotes_from_answer(answer_text)

                            # Remove quotes from answer text to avoid duplication
                            clean_answer_text = answer_text
                            for quote in quotes:
                                clean_answer_text = clean_answer_text.replace(
                                    f'Quote: "{quote}"', ""
                                )
                                clean_answer_text = clean_answer_text.replace(
                                    f"Quote: {quote}", ""
                                )
                                clean_answer_text = clean_answer_text.replace(
                                    f'- Quote: "{quote}"', ""
                                )
                                clean_answer_text = clean_answer_text.replace(
                                    f"- Quote: {quote}", ""
                                )
                                clean_answer_text = clean_answer_text.replace(
                                    f'**Quote:** "{quote}"', ""
                                )
                                clean_answer_text = clean_answer_text.replace(
                                    f"**Quote:** {quote}", ""
                                )

                            # Clean up extra whitespace and empty lines
                            clean_answer_text = re.sub(
                                r"\n\s*\n", "\n", clean_answer_text
                            )
                            clean_answer_text = re.sub(r"^\s*\n", "", clean_answer_text)
                            clean_answer_text = clean_answer_text.strip()

                            # Check if this expert already has an answer for this question
                            existing_answer = None
                            for existing in consolidated["questions_by_section"][
                                section_title
                            ][question]:
                                if existing["citation"] == citation:
                                    existing_answer = existing
                                    break

                            if existing_answer:
                                # Keep the longer/more detailed answer
                                if len(clean_answer_text) > len(
                                    existing_answer["answer"]
                                ):
                                    existing_answer["answer"] = clean_answer_text
                                # Add any new quotes
                                if quotes:
                                    existing_answer["quotes"] = quotes
                            else:
                                # Add new answer
                                consolidated["questions_by_section"][section_title][
                                    question
                                ].append(
                                    {
                                        "answer": clean_answer_text,
                                        "citation": citation,
                                        "quotes": quotes,
                                    }
                                )

        return consolidated

    def export_to_text(self, results: List[Dict[str, Any]], output_file: str = None):
        """Export analysis results to text format with simplified structure."""
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(OUTPUT_FILES["transcript_analysis_txt"])
            else:
                output_file = "FlexXray_transcript_analysis_results.txt"
        try:
            # Consolidate results
            consolidated = self._consolidate_results(results)

            with open(output_file, "w", encoding="utf-8") as f:
                f.write("FLEXXRAY TRANSCRIPT ANALYSIS - SIMPLIFIED STRUCTURE\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"Analysis completed for {len(results)} transcripts.\n")
                f.write(f"Analysis Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")

                # Section 1: Company Overview by Expert
                f.write("SECTION 1: COMPANY & INTERVIEWEE OVERVIEW (BY EXPERT)\n")
                f.write("=" * 60 + "\n\n")

                for citation, overview_data in consolidated[
                    "company_overview_by_expert"
                ].items():
                    f.write(f"Expert: {citation}\n")
                    f.write("-" * 40 + "\n")
                    overview_text = overview_data["answers"].replace("*", "")
                    f.write(f"{overview_text}\n\n")

                # All other sections organized by question
                for section_title, questions_data in consolidated[
                    "questions_by_section"
                ].items():
                    f.write(f"\n{section_title.upper()}\n")
                    f.write("=" * 60 + "\n\n")

                    for question, responses in questions_data.items():
                        f.write(f"Question: {question}\n")
                        f.write("-" * 40 + "\n")

                        if responses:
                            for response in responses:
                                answer_text = response["answer"]
                                citation = response["citation"]
                                quotes = response.get("quotes", [])

                                f.write(f"• {answer_text} (Expert: {citation})\n")

                                # Add quotes if available
                                for quote in quotes:
                                    f.write(f'  "{quote}" -- {citation}\n')
                        else:
                            f.write("• No responses found for this question.\n")

                        f.write("\n")

            print(f"Simplified text file saved to {output_file}")
            return True

        except Exception as e:
            print(f"Error creating text file: {e}")
            return False

    def export_to_word(self, results: List[Dict[str, Any]], output_file: str = None):
        """Export analysis results to Word document format with simplified structure."""
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(OUTPUT_FILES["transcript_analysis_docx"])
            else:
                output_file = "FlexXray_transcript_analysis_results.docx"

        if not DOCX_AVAILABLE:
            print("python-docx not available, saving as text instead...")
            return self.export_to_text(
                results, output_file.replace(".docx", "_fallback.txt")
            )

        try:
            print(f"Creating simplified Word document: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(
                "FlexXray Transcript Analysis - Simplified Structure"
            )
            title_para.alignment = 1
            if title_para.runs:
                title_run = title_para.runs[0]
                title_run.font.size = Inches(0.22)
                title_run.font.bold = True
                title_run.font.name = "Calibri"

            # Add summary section
            summary_heading = doc.add_paragraph("Summary")
            if summary_heading.runs:
                summary_heading_run = summary_heading.runs[0]
                summary_heading_run.font.bold = True
                summary_heading_run.font.size = Inches(0.22)
                summary_heading_run.font.name = "Calibri"

            summary_para = doc.add_paragraph(
                f"Analysis completed for {len(results)} transcripts."
            )
            if summary_para.runs:
                summary_para_run = summary_para.runs[0]
                summary_para_run.font.name = "Calibri"
                summary_para_run.font.size = Inches(0.15)
            summary_para.add_run(
                f'\nAnalysis Date: {time.strftime("%Y-%m-%d %H:%M:%S")}'
            )

            # Section 1: Company Overview by Expert
            doc.add_page_break()
            section1_header = doc.add_paragraph(
                "SECTION 1: COMPANY & INTERVIEWEE OVERVIEW (BY EXPERT)"
            )
            if section1_header.runs:
                section1_header_run = section1_header.runs[0]
                section1_header_run.font.bold = True
                section1_header_run.font.size = Inches(0.24)
                section1_header_run.font.name = "Calibri"

            for citation, overview_data in consolidated[
                "company_overview_by_expert"
            ].items():
                # Add expert header
                expert_header = doc.add_paragraph(f"Expert: {citation}")
                if expert_header.runs:
                    expert_header_run = expert_header.runs[0]
                    expert_header_run.font.bold = True
                    expert_header_run.font.size = Inches(0.20)
                    expert_header_run.font.name = "Calibri"

                # Add the overview answer
                overview_text = overview_data["answers"].replace("*", "")
                overview_para = doc.add_paragraph(overview_text)
                if overview_para.runs:
                    overview_run = overview_para.runs[0]
                    overview_run.font.name = "Calibri"
                    overview_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing between experts

            # All other sections organized by question
            for section_title, questions_data in consolidated[
                "questions_by_section"
            ].items():
                doc.add_page_break()
                section_header = doc.add_paragraph(section_title.upper())
                if section_header.runs:
                    section_header_run = section_header.runs[0]
                    section_header_run.font.bold = True
                    section_header_run.font.size = Inches(0.24)
                    section_header_run.font.name = "Calibri"

                for question, responses in questions_data.items():
                    # Add question with proper formatting
                    question_para = doc.add_paragraph(f"Question: {question}")
                    if question_para.runs:
                        question_run = question_para.runs[0]
                        question_run.font.bold = True
                        question_run.font.size = Inches(0.18)
                        question_run.font.name = "Calibri"

                    if responses:
                        for response in responses:
                            # Add the answer with citation
                            answer_text = response["answer"]
                            citation = response["citation"]
                            quotes = response.get("quotes", [])

                            # Add expert response
                            response_text = f"• {answer_text} (Expert: {citation})"
                            response_para = doc.add_paragraph(response_text)
                            if response_para.runs:
                                response_run = response_para.runs[0]
                                response_run.font.name = "Calibri"
                                response_run.font.size = Inches(0.15)

                            # Add supporting quotes directly under the response
                            for quote in quotes:
                                quote_text = f'  "{quote}" -- {citation}'
                                quote_para = doc.add_paragraph(quote_text)
                                if quote_para.runs:
                                    quote_run = quote_para.runs[0]
                                    quote_run.font.name = "Calibri"
                                    quote_run.font.size = Inches(0.13)
                                    quote_run.font.italic = True
                    else:
                        no_response_para = doc.add_paragraph(
                            "• No responses found for this question."
                        )
                        if no_response_para.runs:
                            no_response_run = no_response_para.runs[0]
                            no_response_run.font.name = "Calibri"
                            no_response_run.font.size = Inches(0.15)

                    # Add spacing between questions
                    doc.add_paragraph("")

            # Save the document
            print(f"Saving simplified Word document to: {output_file}")
            doc.save(output_file)
            print(f"Simplified Word document saved to {output_file}")

            return True

        except Exception as e:
            print(f"Error creating simplified Word document: {e}")
            print("Saving as text instead...")
            return self.export_to_text(
                results, output_file.replace(".docx", "_fallback.txt")
            )

    def create_summarized_analysis(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["transcript_analysis_summary_txt"]
                )
            else:
                output_file = "FlexXray_transcript_analysis_summary.txt"
        """Create a summarized analysis document using OpenAI to identify common themes and key points for each question."""
        try:
            print(f"Creating summarized analysis: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Prepare the summary content
            summary_sections = []

            # Add company overview section
            summary_sections.append("SECTION 1: COMPANY & INTERVIEWEE OVERVIEW")
            summary_sections.append("=" * 50)

            # Summarize company overview by expert
            for citation, overview_data in consolidated[
                "company_overview_by_expert"
            ].items():
                summary_sections.append(f"\nExpert: {citation}")
                summary_sections.append("-" * 30)
                overview_text = overview_data["answers"].replace("*", "")
                summary_sections.append(overview_text)

            # Process each section and question individually
            for section_title, questions_data in consolidated[
                "questions_by_section"
            ].items():
                summary_sections.append(f"\n{section_title.upper()}")
                summary_sections.append("=" * 50)

                for question, responses in questions_data.items():
                    if responses:
                        # Prepare content for this specific question
                        question_content = []
                        question_content.append(f"Question: {question}")
                        question_content.append("-" * 30)

                        for response in responses:
                            answer_text = response["answer"]
                            citation = response["citation"]
                            quotes = response.get("quotes", [])

                            question_content.append(
                                f"• {answer_text} (Expert: {citation})"
                            )

                            # Add quotes if available
                            for quote in quotes:
                                question_content.append(f'  "{quote}" -- {citation}')

                        # Combine all responses for this question
                        question_text = "\n".join(question_content)

                        # Create summarization prompt for this specific question
                        question_prompt = f"""
You are analyzing responses from multiple experts about FlexXray for the following question:

{question}

Below are all the expert responses for this question:

{question_text}

Please identify common themes or recurring points across the expert responses. Summarize those key points in clear, concise bullet points.

Focus on:
1. Common themes and patterns
2. Key insights and findings
3. Important differences or contradictions (include these in the main summary)
4. Actionable takeaways (include these in the main summary)

Please provide a summary with the following structure:

## SUMMARY
[5-8 concise bullet points summarizing the key findings, including any important differences or contradictions between experts, and actionable insights]

Please ensure the summary is clear, concise, and actionable.
"""

                        # Call OpenAI API for this question
                        print(f"Summarizing question: {question[:50]}...")
                        try:
                            response = self.client.chat.completions.create(
                                model="gpt-4",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": "You are an expert business analyst specializing in summarizing and synthesizing information from multiple sources. Focus on identifying patterns, themes, and actionable insights.",
                                    },
                                    {"role": "user", "content": question_prompt},
                                ],
                                temperature=0.3,
                                max_tokens=1000,
                            )

                            question_summary = response.choices[0].message.content

                            # Add the question and its summary to the sections
                            summary_sections.append(f"\nQuestion: {question}")
                            summary_sections.append("-" * 30)
                            summary_sections.append(question_summary)

                        except Exception as e:
                            print(f"Error summarizing question {question[:50]}: {e}")
                            # If summarization fails, just add the raw responses
                            summary_sections.append(f"\nQuestion: {question}")
                            summary_sections.append("-" * 30)
                            summary_sections.append(
                                "Summary generation failed. Raw responses:"
                            )
                            for response in responses:
                                answer_text = response["answer"]
                                citation = response["citation"]
                                summary_sections.append(
                                    f"• {answer_text} (Expert: {citation})"
                                )
                    else:
                        summary_sections.append(f"\nQuestion: {question}")
                        summary_sections.append("-" * 30)
                        summary_sections.append(
                            "• No responses found for this question."
                        )

            # Write the summary to file
            with open(output_file, "w", encoding="utf-8") as f:
                f.write("FLEXXRAY TRANSCRIPT ANALYSIS - AI SUMMARIZED\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"Analysis completed for {len(results)} transcripts.\n")
                f.write(f"Analysis Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Summary generated using OpenAI GPT-4\n\n")
                f.write("=" * 60 + "\n\n")
                f.write("\n".join(summary_sections))

            print(f"Summarized analysis saved to {output_file}")
            return True

        except Exception as e:
            print(f"Error creating summarized analysis: {e}")
            return False

    def create_summarized_analysis_word(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["transcript_analysis_summary_docx"]
                )
            else:
                output_file = "FlexXray_transcript_analysis_summary.docx"
        """Create a summarized analysis Word document using OpenAI to identify common themes and key points for each question."""
        if not DOCX_AVAILABLE:
            print("python-docx not available, saving as text instead...")
            return self.create_summarized_analysis(
                results, output_file.replace(".docx", ".txt")
            )

        try:
            print(f"Creating summarized analysis Word document: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(
                "FlexXray Transcript Analysis - AI Summarized"
            )
            title_para.alignment = 1
            if title_para.runs:
                title_run = title_para.runs[0]
                title_run.font.size = Inches(0.22)
                title_run.font.bold = True
                title_run.font.name = "Calibri"

            # Add summary section
            summary_heading = doc.add_paragraph("Summary")
            if summary_heading.runs:
                summary_heading_run = summary_heading.runs[0]
                summary_heading_run.font.bold = True
                summary_heading_run.font.size = Inches(0.22)
                summary_heading_run.font.name = "Calibri"

            summary_para = doc.add_paragraph(
                f"Analysis completed for {len(results)} transcripts."
            )
            if summary_para.runs:
                summary_para_run = summary_para.runs[0]
                summary_para_run.font.name = "Calibri"
                summary_para_run.font.size = Inches(0.15)
            summary_para.add_run(
                f'\nAnalysis Date: {time.strftime("%Y-%m-%d %H:%M:%S")}'
            )
            summary_para.add_run(f"\nSummary generated using OpenAI GPT-4")

            # Section 1: Company Overview by Expert
            doc.add_page_break()
            section1_header = doc.add_paragraph(
                "SECTION 1: COMPANY & INTERVIEWEE OVERVIEW (BY EXPERT)"
            )
            if section1_header.runs:
                section1_header_run = section1_header.runs[0]
                section1_header_run.font.bold = True
                section1_header_run.font.size = Inches(0.24)
                section1_header_run.font.name = "Calibri"

            # Add company overview by expert
            for citation, overview_data in consolidated[
                "company_overview_by_expert"
            ].items():
                expert_heading = doc.add_paragraph(f"Expert: {citation}")
                if expert_heading.runs:
                    expert_heading_run = expert_heading.runs[0]
                    expert_heading_run.font.bold = True
                    expert_heading_run.font.size = Inches(0.18)
                    expert_heading_run.font.name = "Calibri"

                expert_underline = doc.add_paragraph("-" * 40)
                if expert_underline.runs:
                    expert_underline_run = expert_underline.runs[0]
                    expert_underline_run.font.name = "Calibri"
                    expert_underline_run.font.size = Inches(0.15)

                overview_text = overview_data["answers"].replace("*", "")
                overview_para = doc.add_paragraph(overview_text)
                if overview_para.runs:
                    overview_run = overview_para.runs[0]
                    overview_run.font.name = "Calibri"
                    overview_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing

            # Process each section and question individually
            for section_title, questions_data in consolidated[
                "questions_by_section"
            ].items():
                # Add section header
                doc.add_page_break()
                section_header = doc.add_paragraph(section_title.upper())
                if section_header.runs:
                    section_header_run = section_header.runs[0]
                    section_header_run.font.bold = True
                    section_header_run.font.size = Inches(0.24)
                    section_header_run.font.name = "Calibri"

                section_underline = doc.add_paragraph("=" * 60)
                if section_underline.runs:
                    section_underline_run = section_underline.runs[0]
                    section_underline_run.font.name = "Calibri"
                    section_underline_run.font.size = Inches(0.15)

                for question, responses in questions_data.items():
                    if responses:
                        # Prepare content for this specific question
                        question_content = []
                        question_content.append(f"Question: {question}")
                        question_content.append("-" * 30)

                        for response in responses:
                            answer_text = response["answer"]
                            citation = response["citation"]
                            quotes = response.get("quotes", [])

                            question_content.append(
                                f"• {answer_text} (Expert: {citation})"
                            )

                            # Add quotes if available
                            for quote in quotes:
                                question_content.append(f'  "{quote}" -- {citation}')

                        # Combine all responses for this question
                        question_text = "\n".join(question_content)

                        # Create summarization prompt for this specific question
                        question_prompt = f"""
You are analyzing responses from multiple experts about FlexXray for the following question:

{question}

Below are all the expert responses for this question:

{question_text}

Please identify common themes or recurring points across the expert responses. Summarize those key points in clear, concise bullet points.

Focus on:
1. Common themes and patterns
2. Key insights and findings
3. Important differences or contradictions (include these in the main summary)
4. Actionable takeaways (include these in the main summary)

Please provide a summary with the following structure:

## SUMMARY
[5-8 concise bullet points summarizing the key findings, including any important differences or contradictions between experts, and actionable insights]

Please ensure the summary is clear, concise, and actionable.
"""

                        # Call OpenAI API for this question
                        print(f"Summarizing question for Word doc: {question[:50]}...")
                        try:
                            response = self.client.chat.completions.create(
                                model="gpt-4",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": "You are an expert business analyst specializing in summarizing and synthesizing information from multiple sources. Focus on identifying patterns, themes, and actionable insights.",
                                    },
                                    {"role": "user", "content": question_prompt},
                                ],
                                temperature=0.3,
                                max_tokens=1000,
                            )

                            question_summary = response.choices[0].message.content

                            # Add the question and its summary to the document
                            question_heading = doc.add_paragraph(
                                f"Question: {question}"
                            )
                            if question_heading.runs:
                                question_heading_run = question_heading.runs[0]
                                question_heading_run.font.bold = True
                                question_heading_run.font.size = Inches(0.18)
                                question_heading_run.font.name = "Calibri"

                            question_underline = doc.add_paragraph("-" * 40)
                            if question_underline.runs:
                                question_underline_run = question_underline.runs[0]
                                question_underline_run.font.name = "Calibri"
                                question_underline_run.font.size = Inches(0.15)

                            # Parse and format the summary
                            summary_lines = question_summary.split("\n")
                            for line in summary_lines:
                                if line.strip():
                                    if line.startswith("##"):
                                        # Section header
                                        section_para = doc.add_paragraph(
                                            line.replace("##", "").strip()
                                        )
                                        if section_para.runs:
                                            section_run = section_para.runs[0]
                                            section_run.font.bold = True
                                            section_run.font.size = Inches(0.16)
                                            section_run.font.name = "Calibri"
                                    elif line.startswith("-") or line.startswith("•"):
                                        # Bullet point
                                        bullet_para = doc.add_paragraph(line)
                                        if bullet_para.runs:
                                            bullet_run = bullet_para.runs[0]
                                            bullet_run.font.name = "Calibri"
                                            bullet_run.font.size = Inches(0.15)
                                    else:
                                        # Regular text
                                        text_para = doc.add_paragraph(line)
                                        if text_para.runs:
                                            text_run = text_para.runs[0]
                                            text_run.font.name = "Calibri"
                                            text_run.font.size = Inches(0.15)

                            doc.add_paragraph("")  # Add spacing between questions

                        except Exception as e:
                            print(f"Error summarizing question {question[:50]}: {e}")
                            # If summarization fails, just add the raw responses
                            question_heading = doc.add_paragraph(
                                f"Question: {question}"
                            )
                            if question_heading.runs:
                                question_heading_run = question_heading.runs[0]
                                question_heading_run.font.bold = True
                                question_heading_run.font.size = Inches(0.18)
                                question_heading_run.font.name = "Calibri"

                            question_underline = doc.add_paragraph("-" * 40)
                            if question_underline.runs:
                                question_underline_run = question_underline.runs[0]
                                question_underline_run.font.name = "Calibri"
                                question_underline_run.font.size = Inches(0.15)

                            error_para = doc.add_paragraph(
                                "Summary generation failed. Raw responses:"
                            )
                            if error_para.runs:
                                error_run = error_para.runs[0]
                                error_run.font.name = "Calibri"
                                error_run.font.size = Inches(0.15)

                            for response in responses:
                                answer_text = response["answer"]
                                citation = response["citation"]
                                response_para = doc.add_paragraph(
                                    f"• {answer_text} (Expert: {citation})"
                                )
                                if response_para.runs:
                                    response_run = response_para.runs[0]
                                    response_run.font.name = "Calibri"
                                    response_run.font.size = Inches(0.15)
                    else:
                        question_heading = doc.add_paragraph(f"Question: {question}")
                        if question_heading.runs:
                            question_heading_run = question_heading.runs[0]
                            question_heading_run.font.bold = True
                            question_heading_run.font.size = Inches(0.18)
                            question_heading_run.font.name = "Calibri"

                        question_underline = doc.add_paragraph("-" * 40)
                        if question_underline.runs:
                            question_underline_run = question_underline.runs[0]
                            question_underline_run.font.name = "Calibri"
                            question_underline_run.font.size = Inches(0.15)

                        no_response_para = doc.add_paragraph(
                            "• No responses found for this question."
                        )
                        if no_response_para.runs:
                            no_response_run = no_response_para.runs[0]
                            no_response_run.font.name = "Calibri"
                            no_response_run.font.size = Inches(0.15)

                        doc.add_paragraph("")  # Add spacing

            # Save the document
            print(f"Saving summarized Word document to: {output_file}")
            doc.save(output_file)
            print(f"Summarized Word document saved to {output_file}")

            return True

        except Exception as e:
            print(f"Error creating summarized Word document: {e}")
            print("Saving as text instead...")
            return self.create_summarized_analysis(
                results, output_file.replace(".docx", "_fallback.txt")
            )

    def semantic_search_across_transcripts(
        self, query: str, top_k: int = 10, filter_by_type: str = None
    ) -> List[Dict[str, Any]]:
        """Perform semantic search across all stored transcript chunks."""
        if not self.collection or not CHROMA_AVAILABLE:
            print("Warning: ChromaDB not available. Cannot perform semantic search.")
            return []

        try:
            # Build search query
            search_query = query

            # Add type filter if specified
            where_filter = {}
            if filter_by_type:
                where_filter["chunk_type"] = filter_by_type

            # Search in ChromaDB
            results = self.collection.query(
                query_texts=[search_query],
                n_results=top_k,
                where=where_filter if where_filter else None,
            )

            # Format results
            search_results = []
            if results["documents"] and results["documents"][0]:
                for i, doc in enumerate(results["documents"][0]):
                    result_data = {
                        "text": doc,
                        "metadata": (
                            results["metadatas"][0][i]
                            if results["metadatas"] and results["metadatas"][0]
                            else {}
                        ),
                        "similarity_score": 1
                        - (
                            results["distances"][0][i]
                            if results["distances"] and results["distances"][0]
                            else 0
                        ),
                        "rank": i + 1,
                    }
                    search_results.append(result_data)

            return search_results

        except Exception as e:
            print(f"Error performing semantic search: {e}")
            return []

    def get_chunk_statistics(self) -> Dict[str, Any]:
        """Get statistics about stored chunks in the vector database."""
        if not self.collection or not CHROMA_AVAILABLE:
            return {"error": "ChromaDB not available"}

        try:
            # Get collection count
            count = self.collection.count()

            # Get sample chunks to analyze metadata
            sample_results = self.collection.get(limit=min(100, count))

            # Analyze chunk types
            chunk_types = {}
            transcript_names = set()

            if sample_results["metadatas"]:
                for metadata in sample_results["metadatas"]:
                    if metadata:
                        chunk_type = metadata.get("chunk_type", "unknown")
                        chunk_types[chunk_type] = chunk_types.get(chunk_type, 0) + 1

                        transcript_name = metadata.get("transcript_name", "unknown")
                        transcript_names.add(transcript_name)

            return {
                "total_chunks": count,
                "chunk_types": chunk_types,
                "unique_transcripts": len(transcript_names),
                "transcript_names": list(transcript_names),
                "chunk_size_limit": self.chunk_size,
                "overlap_size": self.overlap_size,
            }

        except Exception as e:
            return {"error": f"Failed to get statistics: {e}"}

    def create_concept_attribution_analysis(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["concept_attribution_analysis_txt"]
                )
            else:
                output_file = "FlexXray_concept_attribution_analysis.txt"
        """Create a detailed analysis showing which experts contributed to each concept/theme."""
        try:
            print(f"Creating concept attribution analysis: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Create concept mapping from vector store
            concept_attribution = self._build_concept_attribution_map()

            with open(output_file, "w", encoding="utf-8") as f:
                f.write(
                    "FLEXXRAY TRANSCRIPT ANALYSIS - CONCEPT ATTRIBUTION BY EXPERT\n"
                )
                f.write("=" * 70 + "\n\n")
                f.write(f"Analysis completed for {len(results)} transcripts.\n")
                f.write(f"Analysis Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(
                    f"Vector database chunks: {len(concept_attribution)} concepts identified\n\n"
                )

                # Section 1: Expert Overview with Key Contributions
                f.write("SECTION 1: EXPERT OVERVIEW & KEY CONTRIBUTIONS\n")
                f.write("=" * 70 + "\n\n")

                expert_contributions = self._analyze_expert_contributions(
                    concept_attribution
                )
                for expert, contributions in expert_contributions.items():
                    f.write(f"EXPERT: {expert}\n")
                    f.write("-" * 50 + "\n")
                    f.write(f"Key Topics: {', '.join(contributions['key_topics'])}\n")
                    f.write(f"Chunks Contributed: {contributions['chunk_count']}\n")
                    f.write(
                        f"Primary Focus Areas: {', '.join(contributions['focus_areas'])}\n"
                    )
                    f.write(
                        f"Notable Insights: {contributions['notable_insights']}\n\n"
                    )

                # Section 2: Concept-by-Concept Attribution
                f.write("SECTION 2: CONCEPT ATTRIBUTION BY BUSINESS THEME\n")
                f.write("=" * 70 + "\n\n")

                # Group concepts by business theme
                themed_concepts = self._group_concepts_by_theme(concept_attribution)

                for theme, concepts in themed_concepts.items():
                    f.write(f"THEME: {theme.upper()}\n")
                    f.write("=" * 50 + "\n\n")

                    for concept, attribution_data in concepts.items():
                        f.write(f"Concept: {concept}\n")
                        f.write("-" * 40 + "\n")

                        # Show expert contributions
                        for expert, details in attribution_data["experts"].items():
                            f.write(f"  • {expert}:\n")
                            f.write(
                                f"    - Relevance Score: {details['relevance']:.3f}\n"
                            )
                            f.write(f"    - Key Quote: \"{details['key_quote']}\"\n")
                            f.write(f"    - Context: {details['context']}\n")

                        f.write(
                            f"  Overall Consensus: {attribution_data['consensus_level']}\n"
                        )
                        f.write(
                            f"  Key Insights: {attribution_data['key_insights']}\n\n"
                        )

                # Section 3: Cross-Expert Analysis
                f.write("SECTION 3: CROSS-EXPERT ANALYSIS & CONSENSUS\n")
                f.write("=" * 70 + "\n\n")

                cross_expert_insights = self._analyze_cross_expert_patterns(
                    concept_attribution
                )
                for insight_type, insights in cross_expert_insights.items():
                    f.write(f"{insight_type.upper()}:\n")
                    f.write("-" * 30 + "\n")
                    for insight in insights:
                        f.write(f"  • {insight}\n")
                    f.write("\n")

                # Section 4: Expert Agreement/Disagreement Matrix
                f.write("SECTION 4: EXPERT AGREEMENT/DISAGREEMENT MATRIX\n")
                f.write("=" * 70 + "\n\n")

                agreement_matrix = self._create_expert_agreement_matrix(
                    concept_attribution
                )
                f.write("Expert Agreement Levels (Higher = More Agreement):\n")
                f.write("-" * 50 + "\n")

                for concept, matrix in agreement_matrix.items():
                    f.write(f"\n{concept}:\n")
                    for expert1 in matrix:
                        for expert2 in matrix[expert1]:
                            if expert1 != expert2:
                                agreement = matrix[expert1][expert2]
                                f.write(f"  {expert1} ↔ {expert2}: {agreement:.2f}\n")

                print(f"Concept attribution analysis saved to {output_file}")
                return True

        except Exception as e:
            print(f"Error creating concept attribution analysis: {e}")
            return False

    def _build_concept_attribution_map(self) -> Dict[str, Any]:
        """Build a comprehensive map of concepts and their expert attributions."""
        if not self.collection or not CHROMA_AVAILABLE:
            return {}

        try:
            # Get all chunks from the collection
            all_chunks = self.collection.get()

            concept_map = {}

            # Process each chunk to extract concepts and expert attributions
            for i, chunk_text in enumerate(all_chunks["documents"]):
                metadata = all_chunks["metadatas"][i] if all_chunks["metadatas"] else {}

                if not metadata:
                    continue

                expert_name = metadata.get("transcript_name", "Unknown")
                chunk_type = metadata.get("chunk_type", "general")

                # Extract key concepts from the chunk
                concepts = self._extract_concepts_from_chunk(chunk_text, chunk_type)

                # Add to concept map
                for concept in concepts:
                    if concept not in concept_map:
                        concept_map[concept] = {
                            "experts": {},
                            "chunk_count": 0,
                            "chunk_type": chunk_type,
                            "total_relevance": 0.0,
                        }

                    # Add expert contribution
                    if expert_name not in concept_map[concept]["experts"]:
                        concept_map[concept]["experts"][expert_name] = {
                            "chunks": [],
                            "relevance": 0.0,
                            "key_quote": "",
                            "context": "",
                        }

                    # Calculate relevance score
                    relevance = self._calculate_concept_relevance(chunk_text, concept)

                    concept_map[concept]["experts"][expert_name]["chunks"].append(
                        {
                            "text": chunk_text,
                            "relevance": relevance,
                            "chunk_id": metadata.get("chunk_id", 0),
                        }
                    )

                    concept_map[concept]["experts"][expert_name]["relevance"] = max(
                        concept_map[concept]["experts"][expert_name]["relevance"],
                        relevance,
                    )

                    # Update key quote if this chunk is more relevant
                    if (
                        relevance
                        > concept_map[concept]["experts"][expert_name]["relevance"]
                        * 0.8
                    ):
                        concept_map[concept]["experts"][expert_name]["key_quote"] = (
                            self._extract_key_quote(chunk_text, concept)
                        )
                        concept_map[concept]["experts"][expert_name]["context"] = (
                            self._extract_context(chunk_text, concept)
                        )

                    concept_map[concept]["chunk_count"] += 1
                    concept_map[concept]["total_relevance"] += relevance

            return concept_map

        except Exception as e:
            print(f"Error building concept attribution map: {e}")
            return {}

    def _extract_concepts_from_chunk(
        self, chunk_text: str, chunk_type: str
    ) -> List[str]:
        """Extract key business concepts from a chunk based on its type."""
        concepts = []
        text_lower = chunk_text.lower()

        # Define concept patterns by chunk type
        concept_patterns = {
            "financial": [
                "revenue growth",
                "profitability",
                "cost structure",
                "pricing strategy",
                "financial performance",
                "investment",
                "funding",
                "cash flow",
            ],
            "competitive_analysis": [
                "competitive advantage",
                "market position",
                "differentiation",
                "strengths",
                "weaknesses",
                "competitive landscape",
            ],
            "risk_challenge": [
                "risk factors",
                "challenges",
                "threats",
                "obstacles",
                "regulatory issues",
                "market risks",
            ],
            "growth_strategy": [
                "growth opportunities",
                "expansion plans",
                "market expansion",
                "product development",
                "strategic initiatives",
            ],
            "customer_market": [
                "customer needs",
                "market demand",
                "customer satisfaction",
                "market trends",
                "customer relationships",
            ],
            "technology_product": [
                "technology innovation",
                "product features",
                "R&D",
                "technical capabilities",
                "product roadmap",
            ],
        }

        # Extract concepts based on chunk type
        if chunk_type in concept_patterns:
            for concept in concept_patterns[chunk_type]:
                if concept in text_lower:
                    concepts.append(concept)

        # Extract FlexXray-specific concepts
        flexxray_concepts = [
            "foreign material detection",
            "x-ray technology",
            "food safety",
            "quality control",
            "contamination prevention",
            "inspection services",
        ]

        for concept in flexxray_concepts:
            if concept in text_lower:
                concepts.append(concept)

        # Extract company names and products
        company_patterns = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", chunk_text)
        for company in company_patterns[:3]:
            if len(company.split()) <= 3 and company.lower() not in [
                "the",
                "and",
                "for",
                "with",
                "flexxray",
            ]:
                concepts.append(f"company: {company}")

        return list(set(concepts))  # Remove duplicates

    def _calculate_concept_relevance(self, chunk_text: str, concept: str) -> float:
        """Calculate how relevant a chunk is to a specific concept."""
        text_lower = chunk_text.lower()
        concept_lower = concept.lower()

        # Simple relevance scoring based on concept frequency and context
        if concept_lower in text_lower:
            # Count occurrences
            occurrences = text_lower.count(concept_lower)

            # Check if concept appears in key positions (beginning, end)
            words = text_lower.split()
            if len(words) > 0:
                if concept_lower in words[:3] or concept_lower in words[-3:]:
                    occurrences += 2  # Bonus for key positions

            # Normalize to 0-1 scale
            relevance = min(1.0, occurrences / 5.0)
            return relevance

        return 0.0

    def _extract_key_quote(self, chunk_text: str, concept: str) -> str:
        """Extract a key quote related to the concept from the chunk."""
        # Look for sentences containing the concept
        sentences = re.split(r"(?<=[.!?])\s+", chunk_text)

        for sentence in sentences:
            if concept.lower() in sentence.lower():
                # Clean up the sentence
                clean_sentence = sentence.strip()
                if len(clean_sentence) > 20 and len(clean_sentence) < 200:
                    return clean_sentence

        # Fallback: return first sentence if no concept-specific sentence found
        if sentences:
            return (
                sentences[0].strip()[:150] + "..."
                if len(sentences[0]) > 150
                else sentences[0].strip()
            )

        return chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text

    def _extract_context(self, chunk_text: str, concept: str) -> str:
        """Extract context around the concept in the chunk."""
        # Find the position of the concept
        concept_pos = chunk_text.lower().find(concept.lower())

        if concept_pos != -1:
            # Get context around the concept (100 characters before and after)
            start = max(0, concept_pos - 100)
            end = min(len(chunk_text), concept_pos + len(concept) + 100)

            context = chunk_text[start:end]

            # Try to start and end at sentence boundaries
            sentences = re.split(r"(?<=[.!?])\s+", context)
            if len(sentences) > 1:
                context = sentences[0] + " " + sentences[-1]

            return context.strip()

        return chunk_text[:150] + "..." if len(chunk_text) > 150 else chunk_text

    def _analyze_expert_contributions(
        self, concept_map: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Analyze what each expert contributed to the overall analysis."""
        expert_analysis = {}

        for concept, data in concept_map.items():
            for expert_name, expert_data in data["experts"].items():
                if expert_name not in expert_analysis:
                    expert_analysis[expert_name] = {
                        "key_topics": set(),
                        "chunk_count": 0,
                        "focus_areas": set(),
                        "notable_insights": [],
                    }

                # Add concept to key topics
                expert_analysis[expert_name]["key_topics"].add(concept)

                # Count chunks
                expert_analysis[expert_name]["chunk_count"] += len(
                    expert_data["chunks"]
                )

                # Add focus areas based on chunk type
                if "chunk_type" in data:
                    expert_analysis[expert_name]["focus_areas"].add(data["chunk_type"])

                # Add notable insights
                if expert_data["relevance"] > 0.7:  # High relevance threshold
                    expert_analysis[expert_name]["notable_insights"].append(concept)

        # Convert sets to lists for output
        for expert in expert_analysis:
            expert_analysis[expert]["key_topics"] = list(
                expert_analysis[expert]["key_topics"]
            )
            expert_analysis[expert]["focus_areas"] = list(
                expert_analysis[expert]["focus_areas"]
            )
            expert_analysis[expert]["notable_insights"] = expert_analysis[expert][
                "notable_insights"
            ][
                :5
            ]  # Top 5

        return expert_analysis

    def _group_concepts_by_theme(
        self, concept_map: Dict[str, Any]
    ) -> Dict[str, Dict[str, Any]]:
        """Group concepts by business theme for better organization."""
        theme_mapping = {
            "financial_performance": [
                "revenue",
                "profit",
                "cost",
                "financial",
                "investment",
            ],
            "competitive_position": [
                "competitive",
                "advantage",
                "strength",
                "weakness",
                "market position",
            ],
            "growth_opportunities": [
                "growth",
                "expansion",
                "opportunity",
                "development",
            ],
            "risk_management": ["risk", "challenge", "threat", "obstacle"],
            "technology_innovation": ["technology", "innovation", "R&D", "technical"],
            "customer_market": ["customer", "market", "demand", "satisfaction"],
            "operational_excellence": ["quality", "efficiency", "process", "operation"],
        }

        themed_concepts = {}

        for concept, data in concept_map.items():
            # Find the best matching theme
            best_theme = "other"
            best_match_score = 0

            for theme, keywords in theme_mapping.items():
                match_score = sum(
                    1 for keyword in keywords if keyword in concept.lower()
                )
                if match_score > best_match_score:
                    best_match_score = match_score
                    best_theme = theme

            if best_theme not in themed_concepts:
                themed_concepts[best_theme] = {}

            # Add consensus and key insights
            data["consensus_level"] = self._calculate_consensus_level(data["experts"])
            data["key_insights"] = self._extract_key_insights(data)

            themed_concepts[best_theme][concept] = data

        return themed_concepts

    def _calculate_consensus_level(self, experts: Dict[str, Any]) -> str:
        """Calculate the level of consensus among experts on a concept."""
        if len(experts) == 1:
            return "Single Expert Opinion"
        elif len(experts) == 2:
            return "Limited Expert Coverage"
        elif len(experts) <= 4:
            return "Moderate Expert Consensus"
        else:
            return "Strong Expert Consensus"

    def _extract_key_insights(self, concept_data: Dict[str, Any]) -> str:
        """Extract key insights from expert contributions to a concept."""
        insights = []

        # Find experts with high relevance
        high_relevance_experts = [
            expert
            for expert, data in concept_data["experts"].items()
            if data["relevance"] > 0.7
        ]

        if high_relevance_experts:
            insights.append(
                f"Strong insights from {len(high_relevance_experts)} experts"
            )

        # Check for agreement/disagreement
        if len(concept_data["experts"]) > 1:
            avg_relevance = sum(
                data["relevance"] for data in concept_data["experts"].values()
            ) / len(concept_data["experts"])
            if avg_relevance > 0.6:
                insights.append("High agreement among experts")
            elif avg_relevance < 0.4:
                insights.append("Mixed opinions among experts")

        return "; ".join(insights) if insights else "Moderate expert coverage"

    def _analyze_cross_expert_patterns(
        self, concept_map: Dict[str, Any]
    ) -> Dict[str, List[str]]:
        """Analyze patterns across experts to identify consensus and disagreements."""
        patterns = {
            "consensus_areas": [],
            "disagreement_areas": [],
            "unique_perspectives": [],
            "complementary_insights": [],
        }

        for concept, data in concept_map.items():
            experts = list(data["experts"].keys())

            if len(experts) == 1:
                patterns["unique_perspectives"].append(
                    f"{concept}: Only mentioned by {experts[0]}"
                )
            elif len(experts) >= 3:
                # Check for consensus
                relevances = [
                    data["experts"][expert]["relevance"] for expert in experts
                ]
                avg_relevance = sum(relevances) / len(relevances)

                if avg_relevance > 0.6:
                    patterns["consensus_areas"].append(
                        f"{concept}: Strong agreement among {len(experts)} experts"
                    )
                elif avg_relevance < 0.4:
                    patterns["disagreement_areas"].append(
                        f"{concept}: Mixed opinions among {len(experts)} experts"
                    )
                else:
                    patterns["complementary_insights"].append(
                        f"{concept}: Complementary perspectives from {len(experts)} experts"
                    )

        return patterns

    def _create_expert_agreement_matrix(
        self, concept_map: Dict[str, Any]
    ) -> Dict[str, Dict[str, Dict[str, float]]]:
        """Create a matrix showing agreement levels between experts on different concepts."""
        agreement_matrix = {}

        for concept, data in concept_map.items():
            experts = list(data["experts"].keys())

            if len(experts) < 2:
                continue

            agreement_matrix[concept] = {}

            for i, expert1 in enumerate(experts):
                agreement_matrix[concept][expert1] = {}

                for j, expert2 in enumerate(experts):
                    if i == j:
                        agreement_matrix[concept][expert1][expert2] = 1.0
                    else:
                        # Calculate agreement based on similarity of their contributions
                        agreement = self._calculate_expert_agreement(
                            data["experts"][expert1], data["experts"][expert2]
                        )
                        agreement_matrix[concept][expert1][expert2] = agreement

        return agreement_matrix

    def _calculate_expert_agreement(
        self, expert1_data: Dict[str, Any], expert2_data: Dict[str, Any]
    ) -> float:
        """Calculate agreement level between two experts on a concept."""
        # Simple agreement calculation based on relevance scores
        relevance1 = expert1_data.get("relevance", 0.0)
        relevance2 = expert2_data.get("relevance", 0.0)

        # If both have high relevance, they likely agree
        if relevance1 > 0.7 and relevance2 > 0.7:
            return 0.9
        elif relevance1 > 0.5 and relevance2 > 0.5:
            return 0.7
        elif relevance1 > 0.3 and relevance2 > 0.3:
            return 0.5
        else:
            return 0.3

    def create_concept_attribution_analysis_word(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["concept_attribution_analysis_docx"]
                )
            else:
                output_file = "FlexXray_concept_attribution_analysis.docx"
        """Create a Word document version of the concept attribution analysis."""
        if not DOCX_AVAILABLE:
            print("python-docx not available, saving as text instead...")
            return self.create_concept_attribution_analysis(
                results, output_file.replace(".docx", ".txt")
            )

        try:
            print(f"Creating concept attribution analysis Word document: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Create concept mapping from vector store
            concept_attribution = self._build_concept_attribution_map()

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(
                "FlexXray Transcript Analysis - Concept Attribution by Expert"
            )
            title_para.alignment = 1
            if title_para.runs:
                title_run = title_para.runs[0]
                title_run.font.size = Inches(0.22)
                title_run.font.bold = True
                title_run.font.name = "Calibri"

            # Add summary section
            summary_heading = doc.add_paragraph("Summary")
            if summary_heading.runs:
                summary_heading_run = summary_heading.runs[0]
                summary_heading_run.font.bold = True
                summary_heading_run.font.size = Inches(0.22)
                summary_heading_run.font.name = "Calibri"

            summary_para = doc.add_paragraph(
                f"Analysis completed for {len(results)} transcripts."
            )
            if summary_para.runs:
                summary_para_run = summary_para.runs[0]
                summary_para_run.font.name = "Calibri"
                summary_para_run.font.size = Inches(0.15)
            summary_para.add_run(
                f'\nAnalysis Date: {time.strftime("%Y-%m-%d %H:%M:%S")}'
            )
            summary_para.add_run(
                f"\nVector database chunks: {len(concept_attribution)} concepts identified"
            )

            # Section 1: Expert Overview with Key Contributions
            doc.add_page_break()
            section1_header = doc.add_paragraph(
                "SECTION 1: EXPERT OVERVIEW & KEY CONTRIBUTIONS"
            )
            if section1_header.runs:
                section1_header_run = section1_header.runs[0]
                section1_header_run.font.bold = True
                section1_header_run.font.size = Inches(0.24)
                section1_header_run.font.name = "Calibri"

            expert_contributions = self._analyze_expert_contributions(
                concept_attribution
            )
            for expert, contributions in expert_contributions.items():
                # Add expert header
                expert_header = doc.add_paragraph(f"EXPERT: {expert}")
                if expert_header.runs:
                    expert_header_run = expert_header.runs[0]
                    expert_header_run.font.bold = True
                    expert_header_run.font.size = Inches(0.20)
                    expert_header_run.font.name = "Calibri"

                # Add expert underline
                expert_underline = doc.add_paragraph("-" * 50)
                if expert_underline.runs:
                    expert_underline_run = expert_underline.runs[0]
                    expert_underline_run.font.name = "Calibri"
                    expert_underline_run.font.size = Inches(0.15)

                # Add contribution details
                details_text = f"Key Topics: {', '.join(contributions['key_topics'])}\n"
                details_text += f"Chunks Contributed: {contributions['chunk_count']}\n"
                details_text += (
                    f"Primary Focus Areas: {', '.join(contributions['focus_areas'])}\n"
                )
                details_text += (
                    f"Notable Insights: {', '.join(contributions['notable_insights'])}"
                )

                details_para = doc.add_paragraph(details_text)
                if details_para.runs:
                    details_run = details_para.runs[0]
                    details_run.font.name = "Calibri"
                    details_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing

            # Section 2: Concept-by-Concept Attribution
            doc.add_page_break()
            section2_header = doc.add_paragraph(
                "SECTION 2: CONCEPT ATTRIBUTION BY BUSINESS THEME"
            )
            if section2_header.runs:
                section2_header_run = section2_header.runs[0]
                section2_header_run.font.bold = True
                section2_header_run.font.size = Inches(0.24)
                section2_header_run.font.name = "Calibri"

            # Group concepts by business theme
            themed_concepts = self._group_concepts_by_theme(concept_attribution)

            for theme, concepts in themed_concepts.items():
                # Add theme header
                theme_header = doc.add_paragraph(f"THEME: {theme.upper()}")
                if theme_header.runs:
                    theme_header_run = theme_header.runs[0]
                    theme_header_run.font.bold = True
                    theme_header_run.font.size = Inches(0.20)
                    theme_header_run.font.name = "Calibri"

                theme_underline = doc.add_paragraph("=" * 50)
                if theme_underline.runs:
                    theme_underline_run = theme_underline.runs[0]
                    theme_underline_run.font.name = "Calibri"
                    theme_underline_run.font.size = Inches(0.15)

                for concept, attribution_data in concepts.items():
                    # Add concept header
                    concept_header = doc.add_paragraph(f"Concept: {concept}")
                    if concept_header.runs:
                        concept_header_run = concept_header.runs[0]
                        concept_header_run.font.bold = True
                        concept_header_run.font.size = Inches(0.18)
                        concept_header_run.font.name = "Calibri"

                    concept_underline = doc.add_paragraph("-" * 40)
                    if concept_underline.runs:
                        concept_underline_run = concept_underline.runs[0]
                        concept_underline_run.font.name = "Calibri"
                        concept_underline_run.font.size = Inches(0.15)

                    # Show expert contributions
                    for expert, details in attribution_data["experts"].items():
                        expert_contribution = doc.add_paragraph(f"  • {expert}:")
                        if expert_contribution.runs:
                            expert_run = expert_contribution.runs[0]
                            expert_run.font.name = "Calibri"
                            expert_run.font.size = Inches(0.15)

                        # Add contribution details
                        detail_lines = [
                            f"    - Relevance Score: {details['relevance']:.3f}",
                            f"    - Key Quote: \"{details['key_quote']}\"",
                            f"    - Context: {details['context']}",
                        ]

                        for line in detail_lines:
                            detail_para = doc.add_paragraph(line)
                            if detail_para.runs:
                                detail_run = detail_para.runs[0]
                                detail_run.font.name = "Calibri"
                                detail_run.font.size = Inches(0.13)

                        doc.add_paragraph("")  # Add spacing between experts

                    # Add overall consensus and key insights
                    consensus_para = doc.add_paragraph(
                        f"  Overall Consensus: {attribution_data['consensus_level']}"
                    )
                    if consensus_para.runs:
                        consensus_run = consensus_para.runs[0]
                        consensus_run.font.name = "Calibri"
                        consensus_run.font.size = Inches(0.15)
                        consensus_run.font.bold = True

                    insights_para = doc.add_paragraph(
                        f"  Key Insights: {attribution_data['key_insights']}"
                    )
                    if insights_para.runs:
                        insights_run = insights_para.runs[0]
                        insights_run.font.name = "Calibri"
                        insights_run.font.size = Inches(0.15)
                        insights_run.font.italic = True

                    doc.add_paragraph("")  # Add spacing between concepts

                doc.add_paragraph("")  # Add spacing between themes

            # Section 3: Cross-Expert Analysis
            doc.add_page_break()
            section3_header = doc.add_paragraph(
                "SECTION 3: CROSS-EXPERT ANALYSIS & CONSENSUS"
            )
            if section3_header.runs:
                section3_header_run = section3_header.runs[0]
                section3_header_run.font.bold = True
                section3_header_run.font.size = Inches(0.24)
                section3_header_run.font.name = "Calibri"

            cross_expert_insights = self._analyze_cross_expert_patterns(
                concept_attribution
            )
            for insight_type, insights in cross_expert_insights.items():
                # Add insight type header
                insight_header = doc.add_paragraph(insight_type.upper())
                if insight_header.runs:
                    insight_header_run = insight_header.runs[0]
                    insight_header_run.font.bold = True
                    insight_header_run.font.size = Inches(0.18)
                    insight_header_run.font.name = "Calibri"

                insight_underline = doc.add_paragraph("-" * 30)
                if insight_underline.runs:
                    insight_underline_run = insight_underline.runs[0]
                    insight_underline_run.font.name = "Calibri"
                    insight_underline_run.font.size = Inches(0.15)

                for insight in insights:
                    insight_para = doc.add_paragraph(f"  • {insight}")
                    if insight_para.runs:
                        insight_run = insight_para.runs[0]
                        insight_run.font.name = "Calibri"
                        insight_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing

            # Section 4: Expert Agreement/Disagreement Matrix
            doc.add_page_break()
            section4_header = doc.add_paragraph(
                "SECTION 4: EXPERT AGREEMENT/DISAGREEMENT MATRIX"
            )
            if section4_header.runs:
                section4_header_run = section4_header.runs[0]
                section4_header_run.font.bold = True
                section4_header_run.font.size = Inches(0.24)
                section4_header_run.font.name = "Calibri"

            agreement_matrix = self._create_expert_agreement_matrix(concept_attribution)

            matrix_header = doc.add_paragraph(
                "Expert Agreement Levels (Higher = More Agreement):"
            )
            if matrix_header.runs:
                matrix_header_run = matrix_header.runs[0]
                matrix_header_run.font.bold = True
                matrix_header_run.font.size = Inches(0.16)
                matrix_header_run.font.name = "Calibri"

            matrix_underline = doc.add_paragraph("-" * 50)
            if matrix_underline.runs:
                matrix_underline_run = matrix_underline.runs[0]
                matrix_underline_run.font.name = "Calibri"
                matrix_underline_run.font.size = Inches(0.15)

            for concept, matrix in agreement_matrix.items():
                # Add concept header
                concept_matrix_header = doc.add_paragraph(f"\n{concept}:")
                if concept_matrix_header.runs:
                    concept_matrix_run = concept_matrix_header.runs[0]
                    concept_matrix_run.font.bold = True
                    concept_matrix_run.font.size = Inches(0.16)
                    concept_matrix_run.font.name = "Calibri"

                for expert1 in matrix:
                    for expert2 in matrix[expert1]:
                        if expert1 != expert2:
                            agreement = matrix[expert1][expert2]
                            agreement_text = f"  {expert1} ↔ {expert2}: {agreement:.2f}"

                            agreement_para = doc.add_paragraph(agreement_text)
                            if agreement_para.runs:
                                agreement_run = agreement_para.runs[0]
                                agreement_run.font.name = "Calibri"
                                agreement_run.font.size = Inches(0.13)

            # Save the document
            print(f"Saving concept attribution Word document to: {output_file}")
            doc.save(output_file)
            print(f"Concept attribution Word document saved to {output_file}")

            return True

        except Exception as e:
            print(f"Error creating concept attribution Word document: {e}")
            print("Saving as text instead...")
            return self.create_concept_attribution_analysis(
                results, output_file.replace(".docx", "_fallback.txt")
            )

    def export_to_text_with_chunks(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["transcript_analysis_with_chunks_txt"]
                )
            else:
                output_file = "FlexXray_transcript_analysis_with_chunks.txt"
        """Export analysis results to text format with enhanced chunk-based insights."""
        try:
            print(f"Creating enhanced text file with chunk insights: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            with open(output_file, "w", encoding="utf-8") as f:
                f.write("FLEXXRAY TRANSCRIPT ANALYSIS - ENHANCED WITH CHUNK INSIGHTS\n")
                f.write("=" * 70 + "\n\n")
                f.write(f"Analysis completed for {len(results)} transcripts.\n")
                f.write(f"Analysis Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")

                # Section 1: Company Overview by Expert (Enhanced with chunks)
                f.write(
                    "SECTION 1: COMPANY & INTERVIEWEE OVERVIEW (BY EXPERT) - ENHANCED\n"
                )
                f.write("=" * 70 + "\n\n")

                for citation, overview_data in consolidated[
                    "company_overview_by_expert"
                ].items():
                    f.write(f"Expert: {citation}\n")
                    f.write("-" * 50 + "\n")

                    # Add overview text
                    overview_text = overview_data["answers"].replace("*", "")
                    f.write(f"Overview: {overview_text}\n\n")

                    # Add chunk-based insights for this expert
                    expert_chunks = self._get_expert_chunks(results, citation)
                    if expert_chunks:
                        f.write("Chunk-Based Insights:\n")
                        f.write("-" * 30 + "\n")

                        # Group chunks by type for better organization
                        chunks_by_type = self._group_chunks_by_type(expert_chunks)
                        for chunk_type, chunks in chunks_by_type.items():
                            f.write(f"\n{chunk_type.replace('_', ' ').title()}:\n")
                            for chunk in chunks[:3]:  # Show top 3 chunks per type
                                f.write(f"  • {chunk['text'][:150]}...\n")

                        f.write("\n")

                # Section 2: Questions by Section (Enhanced with chunk context)
                for section_title, questions_data in consolidated[
                    "questions_by_section"
                ].items():
                    f.write(
                        f"\n{section_title.upper()} - ENHANCED WITH CHUNK CONTEXT\n"
                    )
                    f.write("=" * 70 + "\n\n")

                    for question, responses in questions_data.items():
                        f.write(f"Question: {question}\n")
                        f.write("-" * 50 + "\n")

                        if responses:
                            for response in responses:
                                answer_text = response["answer"]
                                citation = response["citation"]
                                quotes = response.get("quotes", [])

                                f.write(f"• {answer_text} (Expert: {citation})\n")

                                # Add supporting quotes
                                for quote in quotes:
                                    f.write(f'  "{quote}" -- {citation}\n')

                                # Add chunk-based context for this expert and question
                                expert_chunks = self._get_expert_chunks(
                                    results, citation
                                )
                                relevant_chunks = (
                                    self._find_relevant_chunks_for_question(
                                        expert_chunks, question
                                    )
                                )

                                if relevant_chunks:
                                    f.write(f"\n  Supporting Chunk Context:\n")
                                    for chunk in relevant_chunks[
                                        :2
                                    ]:  # Show top 2 relevant chunks
                                        f.write(f"    - {chunk['text'][:120]}...\n")
                                        f.write(
                                            f"      (Chunk Type: {chunk['metadata'].get('chunk_type', 'Unknown')})\n"
                                        )

                                f.write("\n")
                        else:
                            f.write("• No responses found for this question.\n")

                        f.write("\n")

                # Section 3: Chunk Analysis Summary
                f.write("\nSECTION 3: CHUNK ANALYSIS SUMMARY\n")
                f.write("=" * 70 + "\n\n")

                total_chunks = sum(len(result.get("chunks", [])) for result in results)
                f.write(f"Total Chunks Created: {total_chunks}\n")
                f.write(
                    f"Average Chunks per Transcript: {total_chunks / len(results):.1f}\n\n"
                )

                # Chunk type distribution
                chunk_type_distribution = self._analyze_chunk_types(results)
                f.write("Chunk Type Distribution:\n")
                f.write("-" * 30 + "\n")
                for chunk_type, count in chunk_type_distribution.items():
                    percentage = (count / total_chunks) * 100 if total_chunks > 0 else 0
                    f.write(
                        f"  {chunk_type.replace('_', ' ').title()}: {count} chunks ({percentage:.1f}%)\n"
                    )

                f.write("\n")

                # Expert chunk contribution analysis
                f.write("Expert Chunk Contributions:\n")
                f.write("-" * 30 + "\n")
                expert_chunk_counts = self._analyze_expert_chunk_contributions(results)
                for expert, count in expert_chunk_counts.items():
                    f.write(f"  {expert}: {count} chunks\n")

                f.write("\n")

                # Section 4: Semantic Relationships Between Chunks
                f.write("SECTION 4: SEMANTIC RELATIONSHIPS & INSIGHTS\n")
                f.write("=" * 70 + "\n\n")

                # Find common themes across chunks
                common_themes = self._identify_common_themes_across_chunks(results)
                f.write("Common Themes Across Transcripts:\n")
                f.write("-" * 40 + "\n")
                for theme, details in common_themes.items():
                    f.write(f"\nTheme: {theme}\n")
                    f.write(f"  Mentioned by: {', '.join(details['experts'])}\n")
                    f.write(f"  Total chunks: {details['chunk_count']}\n")
                    f.write(f"  Key insight: {details['key_insight']}\n")

                f.write("\n")

                # Cross-expert chunk similarities
                f.write("Cross-Expert Chunk Similarities:\n")
                f.write("-" * 40 + "\n")
                similarities = self._analyze_cross_expert_chunk_similarities(results)
                for expert_pair, similarity_data in similarities.items():
                    f.write(f"\n{expert_pair}:\n")
                    f.write(
                        f"  Similarity Score: {similarity_data['similarity']:.3f}\n"
                    )
                    f.write(
                        f"  Common Topics: {', '.join(similarity_data['common_topics'][:3])}\n"
                    )
                    f.write(
                        f"  Shared Insights: {similarity_data['shared_insights']}\n"
                    )

                f.write("\n")

                # Section 5: Chunk-Based Recommendations
                f.write("SECTION 5: CHUNK-BASED RECOMMENDATIONS\n")
                f.write("=" * 70 + "\n\n")

                recommendations = self._generate_chunk_based_recommendations(results)
                for category, recs in recommendations.items():
                    f.write(f"{category.upper()}:\n")
                    f.write("-" * 30 + "\n")
                    for rec in recs:
                        f.write(f"  • {rec}\n")
                    f.write("\n")

            print(f"Enhanced text file with chunk insights saved to {output_file}")
            return True

        except Exception as e:
            print(f"Error creating enhanced text file: {e}")
            return False

    def _get_expert_chunks(
        self, results: List[Dict[str, Any]], expert_citation: str
    ) -> List[Dict[str, Any]]:
        """Get chunks for a specific expert from the results."""
        for result in results:
            citation = self._extract_citation_enhanced(
                result.get("transcript_name", ""), result.get("doc_path")
            )
            if citation == expert_citation:
                return result.get("chunks", [])
        return []

    def _group_chunks_by_type(
        self, chunks: List[Dict[str, Any]]
    ) -> Dict[str, List[Dict[str, Any]]]:
        """Group chunks by their type for better organization."""
        grouped = {}
        for chunk in chunks:
            chunk_type = chunk.get("metadata", {}).get("chunk_type", "general")
            if chunk_type not in grouped:
                grouped[chunk_type] = []
            grouped[chunk_type].append(chunk)
        return grouped

    def _find_relevant_chunks_for_question(
        self, chunks: List[Dict[str, Any]], question: str
    ) -> List[Dict[str, Any]]:
        """Find chunks that are most relevant to a specific question."""
        if not chunks:
            return []

        # Simple relevance scoring based on question keywords
        question_keywords = set(question.lower().split())

        scored_chunks = []
        for chunk in chunks:
            chunk_text = chunk.get("text", "").lower()
            chunk_keywords = set(chunk_text.split())

            # Calculate overlap
            overlap = len(question_keywords.intersection(chunk_keywords))
            relevance_score = (
                overlap / len(question_keywords) if question_keywords else 0
            )

            scored_chunks.append((chunk, relevance_score))

        # Sort by relevance and return top chunks
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        return [chunk for chunk, score in scored_chunks if score > 0.1][
            :3
        ]  # Top 3 relevant chunks

    def _analyze_chunk_types(self, results: List[Dict[str, Any]]) -> Dict[str, int]:
        """Analyze the distribution of chunk types across all results."""
        chunk_type_counts = {}

        for result in results:
            chunks = result.get("chunks", [])
            for chunk in chunks:
                chunk_type = chunk.get("metadata", {}).get("chunk_type", "general")
                chunk_type_counts[chunk_type] = chunk_type_counts.get(chunk_type, 0) + 1

        return chunk_type_counts

    def _analyze_expert_chunk_contributions(
        self, results: List[Dict[str, Any]]
    ) -> Dict[str, int]:
        """Analyze how many chunks each expert contributed."""
        expert_chunk_counts = {}

        for result in results:
            citation = self._extract_citation_enhanced(
                result.get("transcript_name", ""), result.get("doc_path")
            )
            chunks = result.get("chunks", [])
            expert_chunk_counts[citation] = len(chunks)

        return expert_chunk_counts

    def _identify_common_themes_across_chunks(
        self, results: List[Dict[str, Any]]
    ) -> Dict[str, Dict[str, Any]]:
        """Identify common themes that appear across multiple experts' chunks."""
        theme_data = {}

        for result in results:
            citation = self._extract_citation_enhanced(
                result.get("transcript_name", ""), result.get("doc_path")
            )
            chunks = result.get("chunks", [])

            for chunk in chunks:
                chunk_type = chunk.get("metadata", {}).get("chunk_type", "general")
                chunk_text = chunk.get("text", "")

                if chunk_type not in theme_data:
                    theme_data[chunk_type] = {
                        "experts": set(),
                        "chunk_count": 0,
                        "key_insight": "",
                        "total_relevance": 0.0,
                    }

                theme_data[chunk_type]["experts"].add(citation)
                theme_data[chunk_type]["chunk_count"] += 1

                # Update key insight if this chunk seems more important
                if len(chunk_text) > len(theme_data[chunk_type]["key_insight"]):
                    theme_data[chunk_type]["key_insight"] = (
                        chunk_text[:200] + "..."
                        if len(chunk_text) > 200
                        else chunk_text
                    )

        # Convert sets to lists for output
        for theme in theme_data:
            theme_data[theme]["experts"] = list(theme_data[theme]["experts"])

        return theme_data

    def _analyze_cross_expert_chunk_similarities(
        self, results: List[Dict[str, Any]]
    ) -> Dict[str, Dict[str, Any]]:
        """Analyze similarities between experts based on their chunk content."""
        similarities = {}

        # Get all expert citations
        expert_citations = []
        for result in results:
            citation = self._extract_citation_enhanced(
                result.get("transcript_name", ""), result.get("doc_path")
            )
            if citation not in expert_citations:
                expert_citations.append(citation)

        # Compare each pair of experts
        for i, expert1 in enumerate(expert_citations):
            for j, expert2 in enumerate(expert_citations[i + 1 :], i + 1):
                expert1_chunks = self._get_expert_chunks(results, expert1)
                expert2_chunks = self._get_expert_chunks(results, expert2)

                # Calculate similarity based on chunk types and content
                similarity_score = self._calculate_expert_chunk_similarity(
                    expert1_chunks, expert2_chunks
                )
                common_topics = self._find_common_topics(expert1_chunks, expert2_chunks)
                shared_insights = self._identify_shared_insights(
                    expert1_chunks, expert2_chunks
                )

                pair_name = f"{expert1} ↔ {expert2}"
                similarities[pair_name] = {
                    "similarity": similarity_score,
                    "common_topics": common_topics,
                    "shared_insights": shared_insights,
                }

        return similarities

    def _calculate_expert_chunk_similarity(
        self, chunks1: List[Dict[str, Any]], chunks2: List[Dict[str, Any]]
    ) -> float:
        """Calculate similarity between two sets of chunks."""
        if not chunks1 or not chunks2:
            return 0.0

        # Get chunk types for each expert
        types1 = [
            chunk.get("metadata", {}).get("chunk_type", "general") for chunk in chunks1
        ]
        types2 = [
            chunk.get("metadata", {}).get("chunk_type", "general") for chunk in chunks2
        ]

        # Calculate type overlap
        type_overlap = len(set(types1).intersection(set(types2)))
        total_types = len(set(types1).union(set(types2)))

        if total_types == 0:
            return 0.0

        return type_overlap / total_types

    def _find_common_topics(
        self, chunks1: List[Dict[str, Any]], chunks2: List[Dict[str, Any]]
    ) -> List[str]:
        """Find common topics between two sets of chunks."""
        topics1 = set()
        topics2 = set()

        # Extract topics from chunks
        for chunk in chunks1:
            chunk_type = chunk.get("metadata", {}).get("chunk_type", "general")
            topics1.add(chunk_type)

        for chunk in chunks2:
            chunk_type = chunk.get("metadata", {}).get("chunk_type", "general")
            topics2.add(chunk_type)

        return list(topics1.intersection(topics2))

    def _identify_shared_insights(
        self, chunks1: List[Dict[str, Any]], chunks2: List[Dict[str, Any]]
    ) -> str:
        """Identify shared insights between two sets of chunks."""
        if not chunks1 or not chunks2:
            return "No shared insights"

        # Simple approach: check if they have similar chunk types
        types1 = set(
            chunk.get("metadata", {}).get("chunk_type", "general") for chunk in chunks1
        )
        types2 = set(
            chunk.get("metadata", {}).get("chunk_type", "general") for chunk in chunks2
        )

        common_types = types1.intersection(types2)

        if len(common_types) >= 3:
            return "High overlap in focus areas"
        elif len(common_types) >= 2:
            return "Moderate overlap in focus areas"
        elif len(common_types) >= 1:
            return "Some overlap in focus areas"
        else:
            return "Different focus areas"

    def _generate_chunk_based_recommendations(
        self, results: List[Dict[str, Any]]
    ) -> Dict[str, List[str]]:
        """Generate recommendations based on chunk analysis."""
        recommendations = {
            "focus_areas": [],
            "expert_utilization": [],
            "insight_gaps": [],
            "strategic_opportunities": [],
        }

        # Analyze chunk distribution
        chunk_type_distribution = self._analyze_chunk_types(results)
        total_chunks = sum(chunk_type_distribution.values())

        if total_chunks > 0:
            # Focus areas recommendations
            dominant_types = [
                k for k, v in chunk_type_distribution.items() if v > total_chunks * 0.2
            ]
            if dominant_types:
                recommendations["focus_areas"].append(
                    f"Strong focus on: {', '.join(dominant_types)}"
                )

            underrepresented_types = [
                k for k, v in chunk_type_distribution.items() if v < total_chunks * 0.1
            ]
            if underrepresented_types:
                recommendations["focus_areas"].append(
                    f"Consider expanding coverage of: {', '.join(underrepresented_types)}"
                )

            # Expert utilization recommendations
            expert_chunk_counts = self._analyze_expert_chunk_contributions(results)
            avg_chunks = (
                total_chunks / len(expert_chunk_counts) if expert_chunk_counts else 0
            )

            high_contributors = [
                k for k, v in expert_chunk_counts.items() if v > avg_chunks * 1.5
            ]
            if high_contributors:
                recommendations["expert_utilization"].append(
                    f"High contributors: {', '.join(high_contributors)} - leverage their insights"
                )

            low_contributors = [
                k for k, v in expert_chunk_counts.items() if v < avg_chunks * 0.5
            ]
            if low_contributors:
                recommendations["expert_utilization"].append(
                    f"Low contributors: {', '.join(low_contributors)} - may need additional engagement"
                )

            # Strategic opportunities
            if (
                "growth_strategy" in chunk_type_distribution
                and chunk_type_distribution["growth_strategy"] > total_chunks * 0.15
            ):
                recommendations["strategic_opportunities"].append(
                    "Strong growth focus identified - consider expansion planning"
                )

            if (
                "risk_challenge" in chunk_type_distribution
                and chunk_type_distribution["risk_challenge"] > total_chunks * 0.2
            ):
                recommendations["strategic_opportunities"].append(
                    "Significant risk factors identified - prioritize risk mitigation"
                )

        return recommendations

    def export_to_word_with_chunks(
        self, results: List[Dict[str, Any]], output_file: str = None
    ) -> bool:
        if output_file is None:
            if CONFIG_AVAILABLE:
                output_file = get_output_path(
                    OUTPUT_FILES["transcript_analysis_with_chunks_docx"]
                )
            else:
                output_file = "FlexXray_transcript_analysis_with_chunks.docx"
        """Export analysis results to Word document format with enhanced chunk-based insights."""
        if not DOCX_AVAILABLE:
            print("python-docx not available, saving as text instead...")
            return self.export_to_text_with_chunks(
                results, output_file.replace(".docx", ".txt")
            )

        try:
            print(f"Creating enhanced Word document with chunk insights: {output_file}")

            # Consolidate results
            consolidated = self._consolidate_results(results)

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(
                "FlexXray Transcript Analysis - Enhanced with Chunk Insights"
            )
            title_para.alignment = 1
            if title_para.runs:
                title_run = title_para.runs[0]
                title_run.font.size = Inches(0.22)
                title_run.font.bold = True
                title_run.font.name = "Calibri"

            # Add summary section
            summary_heading = doc.add_paragraph("Summary")
            if summary_heading.runs:
                summary_heading_run = summary_heading.runs[0]
                summary_heading_run.font.bold = True
                summary_heading_run.font.size = Inches(0.22)
                summary_heading_run.font.name = "Calibri"

            summary_para = doc.add_paragraph(
                f"Analysis completed for {len(results)} transcripts."
            )
            if summary_para.runs:
                summary_para_run = summary_para.runs[0]
                summary_para_run.font.name = "Calibri"
                summary_para_run.font.size = Inches(0.15)
            summary_para.add_run(
                f'\nAnalysis Date: {time.strftime("%Y-%m-%d %H:%M:%S")}'
            )

            # Section 1: Company Overview by Expert (Enhanced with chunks)
            doc.add_page_break()
            section1_header = doc.add_paragraph(
                "SECTION 1: COMPANY & INTERVIEWEE OVERVIEW - ENHANCED"
            )
            if section1_header.runs:
                section1_header_run = section1_header.runs[0]
                section1_header_run.font.bold = True
                section1_header_run.font.size = Inches(0.24)
                section1_header_run.font.name = "Calibri"

            for citation, overview_data in consolidated[
                "company_overview_by_expert"
            ].items():
                # Add expert header
                expert_header = doc.add_paragraph(f"Expert: {citation}")
                if expert_header.runs:
                    expert_header_run = expert_header.runs[0]
                    expert_header_run.font.bold = True
                    expert_header_run.font.size = Inches(0.20)
                    expert_header_run.font.name = "Calibri"

                # Add expert underline
                expert_underline = doc.add_paragraph("-" * 50)
                if expert_underline.runs:
                    expert_underline_run = expert_underline.runs[0]
                    expert_underline_run.font.name = "Calibri"
                    expert_underline_run.font.size = Inches(0.15)

                # Add overview text
                overview_text = overview_data["answers"].replace("*", "")
                overview_para = doc.add_paragraph(f"Overview: {overview_text}")
                if overview_para.runs:
                    overview_run = overview_para.runs[0]
                    overview_run.font.name = "Calibri"
                    overview_run.font.size = Inches(0.15)

                # Add chunk-based insights for this expert
                expert_chunks = self._get_expert_chunks(results, citation)
                if expert_chunks:
                    chunk_insights_heading = doc.add_paragraph("Chunk-Based Insights:")
                    if chunk_insights_heading.runs:
                        chunk_insights_run = chunk_insights_heading.runs[0]
                        chunk_insights_run.font.bold = True
                        chunk_insights_run.font.size = Inches(0.16)
                        chunk_insights_run.font.name = "Calibri"

                    chunk_underline = doc.add_paragraph("-" * 30)
                    if chunk_underline.runs:
                        chunk_underline_run = chunk_underline.runs[0]
                        chunk_underline_run.font.name = "Calibri"
                        chunk_underline_run.font.size = Inches(0.15)

                    # Group chunks by type for better organization
                    chunks_by_type = self._group_chunks_by_type(expert_chunks)
                    for chunk_type, chunks in chunks_by_type.items():
                        # Add chunk type header
                        chunk_type_header = doc.add_paragraph(
                            f"{chunk_type.replace('_', ' ').title()}:"
                        )
                        if chunk_type_header.runs:
                            chunk_type_run = chunk_type_header.runs[0]
                            chunk_type_run.font.bold = True
                            chunk_type_run.font.size = Inches(0.15)
                            chunk_type_run.font.name = "Calibri"

                        # Add top 3 chunks for this type
                        for chunk in chunks[:3]:
                            chunk_text = (
                                chunk["text"][:150] + "..."
                                if len(chunk["text"]) > 150
                                else chunk["text"]
                            )
                            chunk_para = doc.add_paragraph(f"  • {chunk_text}")
                            if chunk_para.runs:
                                chunk_run = chunk_para.runs[0]
                                chunk_run.font.name = "Calibri"
                                chunk_run.font.size = Inches(0.13)

                        doc.add_paragraph("")  # Add spacing between chunk types

                    doc.add_paragraph("")  # Add spacing between experts

                doc.add_paragraph("")  # Add spacing between experts

            # Section 2: Questions by Section (Enhanced with chunk context)
            for section_title, questions_data in consolidated[
                "questions_by_section"
            ].items():
                doc.add_page_break()
                section2_header = doc.add_paragraph(
                    f"{section_title.upper()} - ENHANCED WITH CHUNK CONTEXT"
                )
                if section2_header.runs:
                    section2_header_run = section2_header.runs[0]
                    section2_header_run.font.bold = True
                    section2_header_run.font.size = Inches(0.24)
                    section2_header_run.font.name = "Calibri"

                for question, responses in questions_data.items():
                    # Add question header
                    question_header = doc.add_paragraph(f"Question: {question}")
                    if question_header.runs:
                        question_run = question_header.runs[0]
                        question_run.font.bold = True
                        question_run.font.size = Inches(0.18)
                        question_run.font.name = "Calibri"

                    question_underline = doc.add_paragraph("-" * 50)
                    if question_underline.runs:
                        question_underline_run = question_underline.runs[0]
                        question_underline_run.font.name = "Calibri"
                        question_underline_run.font.size = Inches(0.15)

                    if responses:
                        for response in responses:
                            answer_text = response["answer"]
                            citation = response["citation"]
                            quotes = response.get("quotes", [])

                            # Add the answer
                            response_para = doc.add_paragraph(
                                f"• {answer_text} (Expert: {citation})"
                            )
                            if response_para.runs:
                                response_run = response_para.runs[0]
                                response_run.font.name = "Calibri"
                                response_run.font.size = Inches(0.15)

                            # Add supporting quotes
                            for quote in quotes:
                                quote_para = doc.add_paragraph(
                                    f'  "{quote}" -- {citation}'
                                )
                                if quote_para.runs:
                                    quote_run = quote_para.runs[0]
                                    quote_run.font.name = "Calibri"
                                    quote_run.font.size = Inches(0.13)
                                    quote_run.font.italic = True

                            # Add chunk-based context for this expert and question
                            expert_chunks = self._get_expert_chunks(results, citation)
                            relevant_chunks = self._find_relevant_chunks_for_question(
                                expert_chunks, question
                            )

                            if relevant_chunks:
                                chunk_context_heading = doc.add_paragraph(
                                    "  Supporting Chunk Context:"
                                )
                                if chunk_context_heading.runs:
                                    chunk_context_run = chunk_context_heading.runs[0]
                                    chunk_context_run.font.bold = True
                                    chunk_context_run.font.size = Inches(0.13)
                                    chunk_context_run.font.name = "Calibri"

                                for chunk in relevant_chunks[
                                    :2
                                ]:  # Show top 2 relevant chunks
                                    chunk_text = (
                                        chunk["text"][:120] + "..."
                                        if len(chunk["text"]) > 120
                                        else chunk["text"]
                                    )
                                    chunk_context_para = doc.add_paragraph(
                                        f"    - {chunk_text}"
                                    )
                                    if chunk_context_para.runs:
                                        chunk_context_run = chunk_context_para.runs[0]
                                        chunk_context_run.font.name = "Calibri"
                                        chunk_context_run.font.size = Inches(0.12)

                                    chunk_type_para = doc.add_paragraph(
                                        f"      (Chunk Type: {chunk['metadata'].get('chunk_type', 'Unknown')})"
                                    )
                                    if chunk_type_para.runs:
                                        chunk_type_run = chunk_type_para.runs[0]
                                        chunk_type_run.font.name = "Calibri"
                                        chunk_type_run.font.size = Inches(0.11)
                                        chunk_type_run.font.italic = True

                            doc.add_paragraph("")  # Add spacing between responses
                    else:
                        no_response_para = doc.add_paragraph(
                            "• No responses found for this question."
                        )
                        if no_response_para.runs:
                            no_response_run = no_response_para.runs[0]
                            no_response_run.font.name = "Calibri"
                            no_response_run.font.size = Inches(0.15)

                    doc.add_paragraph("")  # Add spacing between questions

                doc.add_paragraph("")  # Add spacing between sections

            # Section 3: Chunk Analysis Summary
            doc.add_page_break()
            section3_header = doc.add_paragraph("SECTION 3: CHUNK ANALYSIS SUMMARY")
            if section3_header.runs:
                section3_header_run = section3_header.runs[0]
                section3_header_run.font.bold = True
                section3_header_run.font.size = Inches(0.24)
                section3_header_run.font.name = "Calibri"

            total_chunks = sum(len(result.get("chunks", [])) for result in results)
            summary_stats_para = doc.add_paragraph(
                f"Total Chunks Created: {total_chunks}"
            )
            if summary_stats_para.runs:
                summary_stats_run = summary_stats_para.runs[0]
                summary_stats_run.font.name = "Calibri"
                summary_stats_run.font.size = Inches(0.15)
                summary_stats_run.font.bold = True

            avg_chunks_para = doc.add_paragraph(
                f"Average Chunks per Transcript: {total_chunks / len(results):.1f}"
            )
            if avg_chunks_para.runs:
                avg_chunks_run = avg_chunks_para.runs[0]
                avg_chunks_run.font.name = "Calibri"
                avg_chunks_run.font.size = Inches(0.15)

            doc.add_paragraph("")  # Add spacing

            # Chunk type distribution
            chunk_type_dist_heading = doc.add_paragraph("Chunk Type Distribution:")
            if chunk_type_dist_heading.runs:
                chunk_type_dist_run = chunk_type_dist_heading.runs[0]
                chunk_type_dist_run.font.bold = True
                chunk_type_dist_run.font.size = Inches(0.16)
                chunk_type_dist_run.font.name = "Calibri"

            chunk_type_underline = doc.add_paragraph("-" * 30)
            if chunk_type_underline.runs:
                chunk_type_underline_run = chunk_type_underline.runs[0]
                chunk_type_underline_run.font.name = "Calibri"
                chunk_type_underline_run.font.size = Inches(0.15)

            chunk_type_distribution = self._analyze_chunk_types(results)
            for chunk_type, count in chunk_type_distribution.items():
                percentage = (count / total_chunks) * 100 if total_chunks > 0 else 0
                chunk_type_para = doc.add_paragraph(
                    f"  {chunk_type.replace('_', ' ').title()}: {count} chunks ({percentage:.1f}%)"
                )
                if chunk_type_para.runs:
                    chunk_type_run = chunk_type_para.runs[0]
                    chunk_type_run.font.name = "Calibri"
                    chunk_type_run.font.size = Inches(0.15)

            doc.add_paragraph("")  # Add spacing

            # Expert chunk contribution analysis
            expert_contrib_heading = doc.add_paragraph("Expert Chunk Contributions:")
            if expert_contrib_heading.runs:
                expert_contrib_run = expert_contrib_heading.runs[0]
                expert_contrib_run.font.bold = True
                expert_contrib_run.font.size = Inches(0.16)
                expert_contrib_run.font.name = "Calibri"

            expert_contrib_underline = doc.add_paragraph("-" * 30)
            if expert_contrib_underline.runs:
                expert_contrib_underline_run = expert_contrib_underline.runs[0]
                expert_contrib_underline_run.font.name = "Calibri"
                expert_contrib_underline_run.font.size = Inches(0.15)

            expert_chunk_counts = self._analyze_expert_chunk_contributions(results)
            for expert, count in expert_chunk_counts.items():
                expert_contrib_para = doc.add_paragraph(f"  {expert}: {count} chunks")
                if expert_contrib_para.runs:
                    expert_contrib_run = expert_contrib_para.runs[0]
                    expert_contrib_run.font.name = "Calibri"
                    expert_contrib_run.font.size = Inches(0.15)

            # Section 4: Semantic Relationships Between Chunks
            doc.add_page_break()
            section4_header = doc.add_paragraph(
                "SECTION 4: SEMANTIC RELATIONSHIPS & INSIGHTS"
            )
            if section4_header.runs:
                section4_header_run = section4_header.runs[0]
                section4_header_run.font.bold = True
                section4_header_run.font.size = Inches(0.24)
                section4_header_run.font.name = "Calibri"

            # Find common themes across chunks
            common_themes_heading = doc.add_paragraph(
                "Common Themes Across Transcripts:"
            )
            if common_themes_heading.runs:
                common_themes_run = common_themes_heading.runs[0]
                common_themes_run.font.bold = True
                common_themes_run.font.size = Inches(0.18)
                common_themes_run.font.name = "Calibri"

            common_themes_underline = doc.add_paragraph("-" * 40)
            if common_themes_underline.runs:
                common_themes_underline_run = common_themes_underline.runs[0]
                common_themes_underline_run.font.name = "Calibri"
                common_themes_underline_run.font.size = Inches(0.15)

            common_themes = self._identify_common_themes_across_chunks(results)
            for theme, details in common_themes.items():
                theme_header = doc.add_paragraph(f"Theme: {theme}")
                if theme_header.runs:
                    theme_run = theme_header.runs[0]
                    theme_run.font.bold = True
                    theme_run.font.size = Inches(0.16)
                    theme_run.font.name = "Calibri"

                theme_details = doc.add_paragraph(
                    f"  Mentioned by: {', '.join(details['experts'])}"
                )
                if theme_details.runs:
                    theme_details_run = theme_details.runs[0]
                    theme_details_run.font.name = "Calibri"
                    theme_details_run.font.size = Inches(0.15)

                theme_chunks = doc.add_paragraph(
                    f"  Total chunks: {details['chunk_count']}"
                )
                if theme_chunks.runs:
                    theme_chunks_run = theme_chunks.runs[0]
                    theme_chunks_run.font.name = "Calibri"
                    theme_chunks_run.font.size = Inches(0.15)

                theme_insight = doc.add_paragraph(
                    f"  Key insight: {details['key_insight']}"
                )
                if theme_insight.runs:
                    theme_insight_run = theme_insight.runs[0]
                    theme_insight_run.font.name = "Calibri"
                    theme_insight_run.font.size = Inches(0.15)
                    theme_insight_run.font.italic = True

                doc.add_paragraph("")  # Add spacing between themes

            # Cross-expert chunk similarities
            cross_expert_heading = doc.add_paragraph("Cross-Expert Chunk Similarities:")
            if cross_expert_heading.runs:
                cross_expert_run = cross_expert_heading.runs[0]
                cross_expert_run.font.bold = True
                cross_expert_run.font.size = Inches(0.18)
                cross_expert_run.font.name = "Calibri"

            cross_expert_underline = doc.add_paragraph("-" * 40)
            if cross_expert_underline.runs:
                cross_expert_underline_run = cross_expert_underline.runs[0]
                cross_expert_underline_run.font.name = "Calibri"
                cross_expert_underline_run.font.size = Inches(0.15)

            similarities = self._analyze_cross_expert_chunk_similarities(results)
            for expert_pair, similarity_data in similarities.items():
                expert_pair_para = doc.add_paragraph(f"{expert_pair}:")
                if expert_pair_para.runs:
                    expert_pair_run = expert_pair_para.runs[0]
                    expert_pair_run.font.bold = True
                    expert_pair_run.font.size = Inches(0.16)
                    expert_pair_run.font.name = "Calibri"

                similarity_score = doc.add_paragraph(
                    f"  Similarity Score: {similarity_data['similarity']:.3f}"
                )
                if similarity_score.runs:
                    similarity_score_run = similarity_score.runs[0]
                    similarity_score_run.font.name = "Calibri"
                    similarity_score_run.font.size = Inches(0.15)

                common_topics = doc.add_paragraph(
                    f"  Common Topics: {', '.join(similarity_data['common_topics'][:3])}"
                )
                if common_topics.runs:
                    common_topics_run = common_topics.runs[0]
                    common_topics_run.font.name = "Calibri"
                    common_topics_run.font.size = Inches(0.15)

                shared_insights = doc.add_paragraph(
                    f"  Shared Insights: {similarity_data['shared_insights']}"
                )
                if shared_insights.runs:
                    shared_insights_run = shared_insights.runs[0]
                    shared_insights_run.font.name = "Calibri"
                    shared_insights_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing between expert pairs

            # Section 5: Chunk-Based Recommendations
            doc.add_page_break()
            section5_header = doc.add_paragraph(
                "SECTION 5: CHUNK-BASED RECOMMENDATIONS"
            )
            if section5_header.runs:
                section5_header_run = section5_header.runs[0]
                section5_header_run.font.bold = True
                section5_header_run.font.size = Inches(0.24)
                section5_header_run.font.name = "Calibri"

            recommendations = self._generate_chunk_based_recommendations(results)
            for category, recs in recommendations.items():
                category_header = doc.add_paragraph(f"{category.upper()}:")
                if category_header.runs:
                    category_header_run = category_header.runs[0]
                    category_header_run.font.bold = True
                    category_header_run.font.size = Inches(0.18)
                    category_header_run.font.name = "Calibri"

                category_underline = doc.add_paragraph("-" * 30)
                if category_underline.runs:
                    category_underline_run = category_underline.runs[0]
                    category_underline_run.font.name = "Calibri"
                    category_underline_run.font.size = Inches(0.15)

                for rec in recs:
                    rec_para = doc.add_paragraph(f"  • {rec}")
                    if rec_para.runs:
                        rec_run = rec_para.runs[0]
                        rec_run.font.name = "Calibri"
                        rec_run.font.size = Inches(0.15)

                doc.add_paragraph("")  # Add spacing between categories

            # Save the document
            print(
                f"Saving enhanced Word document with chunk insights to: {output_file}"
            )
            doc.save(output_file)
            print(f"Enhanced Word document with chunk insights saved to {output_file}")

            return True

        except Exception as e:
            print(f"Error creating enhanced Word document with chunk insights: {e}")
            print("Saving as text instead...")
            return self.export_to_text_with_chunks(
                results, output_file.replace(".docx", "_fallback.txt")
            )

    def hierarchical_retrieval_with_reranking(
        self,
        query: str,
        top_k_per_transcript: int = 5,
        final_top_k: int = 10,
        use_llm_reranking: bool = True,
    ) -> List[Dict[str, Any]]:
        """
        Hierarchical retrieval: get top N chunks per transcript, then rerank using
        embedding similarity + LLM scoring for better relevance.

        Args:
            query: Search query
            top_k_per_transcript: Number of top chunks to retrieve per transcript
            final_top_k: Final number of top chunks after reranking
            use_llm_reranking: Whether to use LLM scoring for reranking

        Returns:
            List of reranked chunks with scores and metadata
        """
        if not self.collection or not CHROMA_AVAILABLE:
            print(
                "Warning: ChromaDB not available. Cannot perform hierarchical retrieval."
            )
            return []

        try:
            print(f"Performing hierarchical retrieval for query: '{query}'")
            print(
                f"Retrieving top {top_k_per_transcript} chunks per transcript, then reranking to top {final_top_k}"
            )

            # Step 1: Get all transcripts
            all_chunks = self.collection.get()
            transcript_names = set()
            for metadata in all_chunks["metadatas"]:
                if metadata and "transcript_name" in metadata:
                    transcript_names.add(metadata["transcript_name"])

            print(
                f"Found {len(transcript_names)} transcripts: {', '.join(transcript_names)}"
            )

            # Step 2: Retrieve top N chunks per transcript using embedding similarity
            transcript_chunks = {}
            for transcript_name in transcript_names:
                # Search within this specific transcript
                results = self.collection.query(
                    query_texts=[query],
                    n_results=top_k_per_transcript,
                    where={"transcript_name": transcript_name},
                )

                if results["documents"] and results["documents"][0]:
                    transcript_chunks[transcript_name] = []
                    for i, doc in enumerate(results["documents"][0]):
                        chunk_data = {
                            "text": doc,
                            "metadata": (
                                results["metadatas"][0][i]
                                if results["metadatas"] and results["metadatas"][0]
                                else {}
                            ),
                            "embedding_score": 1.0
                            - (
                                results["distances"][0][i]
                                if results["distances"] and results["distances"][0]
                                else 0.0
                            ),
                            "transcript_name": transcript_name,
                        }
                        transcript_chunks[transcript_name].append(chunk_data)

            # Step 3: Combine all chunks and calculate initial scores
            all_candidate_chunks = []
            for transcript_name, chunks in transcript_chunks.items():
                print(f"  {transcript_name}: Retrieved {len(chunks)} chunks")
                all_candidate_chunks.extend(chunks)

            if not all_candidate_chunks:
                print("No chunks found for any transcript")
                return []

            print(f"Total candidate chunks: {len(all_candidate_chunks)}")

            # Step 4: Rerank chunks using LLM scoring if enabled
            if use_llm_reranking and all_candidate_chunks:
                print("Applying LLM-based reranking...")
                reranked_chunks = self._llm_rerank_chunks(
                    query, all_candidate_chunks, final_top_k
                )
            else:
                # Simple reranking based on embedding scores
                print("Using embedding-based reranking...")
                reranked_chunks = self._embedding_based_rerank(
                    all_candidate_chunks, final_top_k
                )

            # Step 5: Format final results
            final_results = []
            for i, chunk in enumerate(reranked_chunks):
                # Calculate final score based on available scoring methods
                if "llm_score" in chunk and "combined_score" in chunk:
                    # LLM reranking was successful
                    final_score = chunk["combined_score"]
                    llm_score = chunk["llm_score"]
                else:
                    # Fallback to embedding-based scoring
                    final_score = chunk["embedding_score"]
                    llm_score = chunk.get("llm_score", "N/A")

                result = {
                    "rank": i + 1,
                    "text": chunk["text"],
                    "metadata": chunk["metadata"],
                    "transcript_name": chunk["transcript_name"],
                    "embedding_score": chunk["embedding_score"],
                    "final_score": final_score,
                    "llm_score": llm_score,
                    "chunk_type": chunk["metadata"].get("chunk_type", "unknown"),
                    "chunk_id": chunk["metadata"].get("chunk_id", 0),
                }
                final_results.append(result)

            print(
                f"Hierarchical retrieval complete. Returning top {len(final_results)} results."
            )
            return final_results

        except Exception as e:
            print(f"Error in hierarchical retrieval: {e}")
            return []

    def _llm_rerank_chunks(
        self, query: str, chunks: List[Dict[str, Any]], final_top_k: int
    ) -> List[Dict[str, Any]]:
        """
        Rerank chunks using LLM scoring for relevance to the query.
        """
        try:
            # Prepare chunks for LLM scoring
            chunk_texts = []
            for chunk in chunks:
                # Truncate text if too long for LLM
                chunk_text = (
                    chunk["text"][:1000] if len(chunk["text"]) > 1000 else chunk["text"]
                )
                chunk_texts.append(chunk_text)

            # Create scoring prompt
            scoring_prompt = f"""
            You are an expert analyst evaluating the relevance of text chunks to a specific query.
            
            QUERY: {query}
            
            Please score each of the following text chunks on a scale of 0.0 to 1.0 based on their relevance to the query.
            Consider:
            - Direct answer to the query
            - Supporting information
            - Contextual relevance
            - Specificity of information
            
            Return ONLY a JSON array of scores, one for each chunk, in order.
            Example: [0.8, 0.3, 0.9, 0.6]
            
            TEXT CHUNKS:
            """

            for i, text in enumerate(chunk_texts):
                scoring_prompt += f"\n--- CHUNK {i+1} ---\n{text}\n"

            # Get LLM scores
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a precise scoring system. Return only valid JSON arrays.",
                    },
                    {"role": "user", "content": scoring_prompt},
                ],
                temperature=0.1,
                max_tokens=500,
            )

            # Parse LLM response
            llm_response = response.choices[0].message.content.strip()
            print(f"LLM Response: {llm_response[:200]}...")

            # Extract scores (handle various response formats)
            scores = self._extract_scores_from_llm_response(llm_response, len(chunks))
            print(f"Extracted scores: {scores}")

            if not scores:
                print(
                    "Warning: Could not parse LLM scores, falling back to embedding-based reranking"
                )
                return self._embedding_based_rerank(chunks, final_top_k)

            # Add LLM scores to chunks
            for i, chunk in enumerate(chunks):
                chunk["llm_score"] = scores[i] if i < len(scores) else 0.0

            # Rerank by combined score (embedding + LLM)
            for chunk in chunks:
                # Combine scores: 40% embedding, 60% LLM
                embedding_weight = 0.4
                llm_weight = 0.6
                chunk["combined_score"] = (
                    embedding_weight * chunk["embedding_score"]
                    + llm_weight * chunk["llm_score"]
                )

            # Sort by combined score and return top K
            reranked_chunks = sorted(
                chunks, key=lambda x: x["combined_score"], reverse=True
            )
            return reranked_chunks[:final_top_k]

        except Exception as e:
            print(f"Error in LLM reranking: {e}")
            print("Falling back to embedding-based reranking")
            return self._embedding_based_rerank(chunks, final_top_k)

    def _extract_scores_from_llm_response(
        self, response: str, expected_count: int
    ) -> List[float]:
        """
        Extract numerical scores from LLM response, handling various formats.
        """
        try:
            # Try to find JSON array
            import json
            import re

            # Look for JSON array pattern
            json_match = re.search(r"\[[\d\.,\s]+\]", response)
            if json_match:
                scores = json.loads(json_match.group())
                if isinstance(scores, list) and len(scores) == expected_count:
                    return [float(score) for score in scores]

            # Look for comma-separated numbers
            numbers = re.findall(r"\b\d+\.?\d*\b", response)
            if len(numbers) >= expected_count:
                return [float(num) for num in numbers[:expected_count]]

            # Look for "score: X" patterns
            score_pattern = re.findall(
                r"score[:\s]*(\d+\.?\d*)", response, re.IGNORECASE
            )
            if len(score_pattern) >= expected_count:
                return [float(score) for score in score_pattern[:expected_count]]

            print(f"Could not extract scores from LLM response: {response}")
            return []

        except Exception as e:
            print(f"Error extracting scores: {e}")
            return []

    def _embedding_based_rerank(
        self, chunks: List[Dict[str, Any]], final_top_k: int
    ) -> List[Dict[str, Any]]:
        """
        Simple reranking based on embedding scores only.
        """
        # Sort by embedding score and return top K
        reranked_chunks = sorted(
            chunks, key=lambda x: x["embedding_score"], reverse=True
        )
        return reranked_chunks[:final_top_k]

    def enhanced_context_retrieval(
        self,
        query: str,
        transcript_name: str = None,
        use_hierarchical: bool = True,
        top_k_per_transcript: int = 5,
        final_top_k: int = 10,
    ) -> str:
        """
        Enhanced context retrieval using hierarchical retrieval and reranking.
        """
        if use_hierarchical:
            # Use hierarchical retrieval with reranking
            relevant_chunks = self.hierarchical_retrieval_with_reranking(
                query,
                top_k_per_transcript=top_k_per_transcript,
                final_top_k=final_top_k,
            )
        else:
            # Fall back to original simple search
            relevant_chunks = self.search_relevant_chunks(
                query, transcript_name, final_top_k
            )

        if not relevant_chunks:
            return ""

        # Format context with enhanced information
        context_parts = []
        for chunk in relevant_chunks:
            if isinstance(chunk, dict) and "text" in chunk:
                # Enhanced format with scores and metadata
                if "rank" in chunk:
                    # Hierarchical retrieval result
                    context_parts.append(
                        f"Rank {chunk['rank']} (Score: {chunk.get('final_score', 'N/A'):.3f}) "
                        f"from {chunk.get('transcript_name', 'Unknown')} "
                        f"[{chunk.get('chunk_type', 'unknown')}]:\n{chunk['text']}"
                    )
                else:
                    # Simple search result
                    context_parts.append(
                        f"Context from {chunk.get('metadata', {}).get('transcript_name', 'Unknown')}:\n{chunk['text']}"
                    )

        return "\n\n".join(context_parts)

    def extract_and_select_best_quotes(
        self,
        question: str,
        results: List[Dict],
        max_quotes_per_expert: int = 4,
        max_final_quotes: int = 3,
        debug: bool = False,
    ) -> List[Dict]:
        """
        Extract top quotes from each expert for a specific question, then use LLM to select the best 1-2 quotes.

        Args:
            question: The specific question to find quotes for
            results: List of transcript analysis results
            max_quotes_per_expert: Maximum quotes to extract per expert
            max_final_quotes: Maximum final quotes to select across all experts

        Returns:
            List of selected quote dictionaries with metadata
        """
        try:
            all_expert_quotes = []

            # Extract quotes from each expert's transcript
            for result in results:
                transcript_name = result.get("transcript_name", "Unknown")
                transcript_text = result.get("transcript_text", "")

                # Search for relevant chunks containing quotes for this question
                relevant_chunks = self.search_relevant_chunks(
                    question,
                    transcript_name,
                    top_k=max_quotes_per_expert
                    * 3,  # Get more chunks to find better quotes
                )

                # Extract potential quotes from chunks
                for chunk in relevant_chunks:
                    chunk_text = chunk.get("text", "")

                    # Look for quote-like patterns (text in quotes, statements, etc.)
                    potential_quotes = self._extract_potential_quotes(
                        chunk_text, question
                    )

                    for quote in potential_quotes:
                        if (
                            len(quote) > 15 and len(quote) < 400
                        ):  # More lenient quote length
                            # Calculate quote relevance to the question
                            quote_relevance = self._calculate_quote_relevance(
                                quote, question
                            )
                            chunk_score = chunk.get("score", 0.0)

                            # Ensure we have a valid relevance score
                            if chunk_score == 0.0:
                                chunk_score = 0.5  # Default chunk score

                            final_relevance = chunk_score * quote_relevance

                            all_expert_quotes.append(
                                {
                                    "quote": quote,
                                    "expert": transcript_name,
                                    "chunk_type": chunk.get("metadata", {}).get(
                                        "chunk_type", "general"
                                    ),
                                    "relevance_score": final_relevance,
                                    "source_chunk": (
                                        chunk_text[:100] + "..."
                                        if len(chunk_text) > 100
                                        else chunk_text
                                    ),
                                }
                            )

            if debug:
                print(f"Found {len(all_expert_quotes)} total quotes across all experts")
                print(
                    f"Sample quotes: {[q['quote'][:50] + '...' for q in all_expert_quotes[:3]]}"
                )

            if not all_expert_quotes:
                return []

            # Use LLM to select the best quotes
            selected_quotes = self._llm_select_best_quotes(
                question, all_expert_quotes, max_final_quotes
            )

            if debug:
                print(f"LLM selected {len(selected_quotes)} quotes")

            return selected_quotes

        except Exception as e:
            print(f"Error extracting and selecting quotes: {e}")
            return []

    def _extract_potential_quotes(self, text: str, question: str) -> List[str]:
        """Extract potential quotes from chunk text."""
        quotes = []

        # Look for quoted text
        import re

        quoted_pattern = r'"([^"]{20,300})"'
        quoted_matches = re.findall(quoted_pattern, text)
        quotes.extend(quoted_matches)

        # Look for statements that might be quotes (sentences that could be direct speech)
        sentences = re.split(r"[.!?]+", text)
        for sentence in sentences:
            sentence = sentence.strip()
            # More lenient criteria for what constitutes a quote
            if (
                len(sentence) > 15  # Reduced minimum length
                and len(sentence) < 400  # Increased maximum length
                and not sentence.startswith("Interviewer:")
                and not sentence.startswith("Interviewee:")
                and not sentence.startswith("Paul Rodriguez")  # Skip interviewer names
                and not sentence.startswith("Gaurav")
                and not sentence.startswith("Randy Jesberg")
                and not sentence.startswith("Cheryl Bertics")
                and not sentence.startswith("Eli Dantas")
                and not sentence.startswith("George Perry")
                and not sentence.startswith("Lee Reece")
                and not sentence.startswith("Peter Poteres")
                and
                # Look for meaningful content indicators
                any(
                    keyword in sentence.lower()
                    for keyword in [
                        "flexxray",
                        "company",
                        "business",
                        "market",
                        "customer",
                        "technology",
                        "service",
                        "quality",
                        "safety",
                        "detection",
                        "x-ray",
                        "inspection",
                        "competitive",
                        "advantage",
                        "growth",
                        "opportunity",
                        "risk",
                        "challenge",
                        "industry",
                        "trend",
                        "future",
                        "strategy",
                        "operation",
                        "process",
                    ]
                )
            ):
                quotes.append(sentence)

        # Also look for key insights in shorter phrases
        insight_patterns = [
            r"FlexXray\s+[^.!?]{10,100}[.!?]",  # FlexXray-related insights
            r"[A-Z][^.!?]{20,150}\s+(?:is|are|has|have|provides|offers|delivers)[^.!?]{10,100}[.!?]",  # Statements about capabilities
            r"[A-Z][^.!?]{20,150}\s+(?:advantage|strength|benefit|value|quality)[^.!?]{10,100}[.!?]",  # Positive attributes
        ]

        for pattern in insight_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if len(match) > 20 and len(match) < 300:
                    quotes.append(match.strip())

        # Remove duplicates and limit results
        unique_quotes = list(
            dict.fromkeys(quotes)
        )  # Preserve order while removing duplicates
        return unique_quotes[:8]  # Increased limit to 8 potential quotes per chunk

    def _calculate_quote_relevance(self, quote: str, question: str) -> float:
        """Calculate how relevant a quote is to a specific question."""
        try:
            # Convert to lowercase for comparison
            quote_lower = quote.lower()
            question_lower = question.lower()

            # Define relevance keywords for different question types
            relevance_keywords = {
                "competitive": [
                    "competitive",
                    "advantage",
                    "strength",
                    "benefit",
                    "superior",
                    "better",
                    "leader",
                    "market leader",
                ],
                "growth": [
                    "growth",
                    "opportunity",
                    "expand",
                    "market",
                    "future",
                    "potential",
                    "increase",
                    "development",
                ],
                "risk": [
                    "risk",
                    "challenge",
                    "problem",
                    "issue",
                    "concern",
                    "threat",
                    "obstacle",
                    "difficulty",
                ],
                "technology": [
                    "technology",
                    "x-ray",
                    "detection",
                    "equipment",
                    "innovation",
                    "capability",
                    "system",
                ],
                "customer": [
                    "customer",
                    "client",
                    "service",
                    "quality",
                    "satisfaction",
                    "experience",
                    "needs",
                ],
                "financial": [
                    "financial",
                    "revenue",
                    "cost",
                    "pricing",
                    "profit",
                    "investment",
                    "return",
                    "value",
                ],
            }

            # Determine question type and get relevant keywords
            question_type = "general"
            for key in relevance_keywords:
                if key in question_lower:
                    question_type = key
                    break

            # Calculate relevance score
            relevant_keywords = relevance_keywords.get(question_type, [])
            keyword_matches = sum(
                1 for keyword in relevant_keywords if keyword in quote_lower
            )

            # Base relevance score
            relevance_score = 0.5  # Default score

            # Boost score based on keyword matches
            if keyword_matches > 0:
                relevance_score += (keyword_matches / len(relevant_keywords)) * 0.5

            # Boost score if quote contains key entities
            if "flexxray" in quote_lower:
                relevance_score += 0.2

            # Boost score for quotes that seem like direct statements
            if quote.startswith(("FlexXray", "The company", "We", "Our", "It")):
                relevance_score += 0.1

            return min(relevance_score, 1.0)  # Cap at 1.0

        except Exception as e:
            print(f"Error calculating quote relevance: {e}")
            return 0.5  # Default relevance score

    def _llm_select_best_quotes(
        self, question: str, candidate_quotes: List[Dict], max_quotes: int
    ) -> List[Dict]:
        """Use LLM to select the best quotes for a given question."""
        try:
            if not candidate_quotes:
                return []

            # Prepare quote candidates for LLM
            quote_candidates = []
            for i, quote_data in enumerate(candidate_quotes):
                quote_candidates.append(
                    f"""
Quote {i+1}:
- Text: "{quote_data['quote']}"
- Expert: {quote_data['expert']}
- Relevance Score: {quote_data['relevance_score']:.3f}
- Chunk Type: {quote_data['chunk_type']}
"""
                )

            quote_candidates_text = "\n".join(quote_candidates)

            prompt = f"""
You are selecting the best quotes to answer a specific question from a business analysis.

QUESTION: {question}

AVAILABLE QUOTES:
{quote_candidates_text}

TASK: Select the {max_quotes} most relevant and insightful quotes that best answer the question.

SELECTION CRITERIA:
1. **Relevance**: Quote directly addresses the question
2. **Insight Value**: Quote provides unique or valuable information
3. **Expert Credibility**: Consider the expert's role and company
4. **Clarity**: Quote is clear and understandable
5. **Diversity**: If possible, select quotes from different experts

INSTRUCTIONS:
- Return ONLY the quote numbers (e.g., "1, 3" or "2")
- Select exactly {max_quotes} quotes
- Choose quotes that together provide the most comprehensive answer
- If no quotes are truly relevant, return "None"

RESPONSE FORMAT: Just the quote numbers separated by commas, or "None"
"""

            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a quote selection expert. Return only the quote numbers or 'None'.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=50,
            )

            response_text = response.choices[0].message.content.strip()

            # Parse the response to get selected quote numbers
            if response_text.lower() == "none":
                return []

            try:
                # Extract numbers from response
                import re

                numbers = [int(n.strip()) for n in re.findall(r"\d+", response_text)]

                # Select the quotes based on LLM selection
                selected_quotes = []
                for num in numbers[:max_quotes]:
                    if 1 <= num <= len(candidate_quotes):
                        selected_quotes.append(candidate_quotes[num - 1])

                return selected_quotes

            except (ValueError, IndexError):
                # Fallback: return top quotes by relevance score
                sorted_quotes = sorted(
                    candidate_quotes, key=lambda x: x["relevance_score"], reverse=True
                )
                return sorted_quotes[:max_quotes]

        except Exception as e:
            print(f"Error in LLM quote selection: {e}")
            print("Falling back to relevance-based selection...")
            # Fallback: return top quotes by relevance score
            try:
                sorted_quotes = sorted(
                    candidate_quotes, key=lambda x: x["relevance_score"], reverse=True
                )
                return sorted_quotes[:max_quotes]
            except Exception as fallback_error:
                print(f"Fallback selection also failed: {fallback_error}")
                # Last resort: return first few quotes
                return candidate_quotes[:max_quotes] if candidate_quotes else []


def main():
    """Main function to run the transcript summarizer."""
    print("Transcript Summarizer with Semantic Chunking and Vector Storage")
    print("=" * 60)

    # Get API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        api_key = input("Please enter your OpenAI API key: ").strip()

    try:
        # Initialize summarizer with ChromaDB
        chroma_dir = "./chroma_db"
        summarizer = TranscriptSummarizer(
            api_key, max_workers=3, chroma_persist_directory=chroma_dir
        )

        # Set default directory to "FlexXray Transcripts"
        default_directory = "FlexXray Transcripts"

        # Check if FlexXray Transcripts folder exists
        if os.path.exists(default_directory):
            directory_path = default_directory
            print(f"Using default directory: {directory_path}")
        else:
            # If FlexXray Transcripts doesn't exist, ask user for directory
            print(f"Default directory '{default_directory}' not found.")
            directory_path = input(
                "Enter the path to the directory containing Word document transcripts: "
            ).strip()

        if not os.path.exists(directory_path):
            print(f"Directory {directory_path} does not exist.")
            return

        # Process transcripts with progress tracking
        print(f"\nProcessing transcripts in: {directory_path}")
        start_time = time.time()

        def progress_callback(message, progress):
            print(f"{message} - {progress:.1f}%")

        results = summarizer.process_transcripts_directory(
            directory_path, progress_callback
        )

        # Store results in the class instance for quote extraction
        summarizer.results = results

        end_time = time.time()
        processing_time = end_time - start_time

        if results:
            # Save results
            summarizer.save_results(results)
            summarizer.export_to_text(results)
            summarizer.export_to_word(results)

            # Create enhanced export with chunk insights
            summarizer.export_to_text_with_chunks(results)
            summarizer.export_to_word_with_chunks(results)

            # Create summarized analysis
            summarizer.create_summarized_analysis(results)
            summarizer.create_summarized_analysis_word(results)

            # Create concept attribution analysis
            summarizer.create_concept_attribution_analysis(results)
            summarizer.create_concept_attribution_analysis_word(results)

            print(
                f"\nAnalysis complete! Processed {len(results)} transcripts in {processing_time:.2f} seconds."
            )
            print("Files created:")
            print("- FlexXray_transcript_analysis_results.json (structured data)")
            print("- FlexXray_transcript_analysis_results.txt (simplified text report)")
            print(
                "- FlexXray_transcript_analysis_results.docx (simplified Word document)"
            )
            print(
                "- FlexXray_transcript_analysis_with_chunks.txt (enhanced with chunk insights)"
            )
            print(
                "- FlexXray_transcript_analysis_with_chunks.docx (enhanced Word document with chunk insights)"
            )
            print("- FlexXray_transcript_analysis_summary.txt (AI-summarized analysis)")
            print(
                "- FlexXray_transcript_analysis_summary.docx (AI-summarized Word document)"
            )
            print("- FlexXray_concept_attribution_analysis.txt (concept attribution)")
            print(
                "- FlexXray_concept_attribution_analysis.docx (concept attribution Word document)"
            )

            # Show vector database statistics
            if summarizer.collection:
                print("\nVector Database Statistics:")
                stats = summarizer.get_chunk_statistics()
                if "error" not in stats:
                    print(f"- Total chunks stored: {stats['total_chunks']}")
                    print(f"- Unique transcripts: {stats['unique_transcripts']}")
                    print(
                        f"- Chunk types: {', '.join([f'{k}: {v}' for k, v in stats['chunk_types'].items()])}"
                    )
                    print(
                        f"- Chunk size: {stats['chunk_size_limit']} chars with {stats['overlap_size']} char overlap"
                    )

                # Demonstrate semantic search capabilities
                print("\nDemonstrating semantic search capabilities...")
                demo_queries = [
                    "What are FlexXray's competitive advantages?",
                    "What growth opportunities exist?",
                    "What are the main risks mentioned?",
                ]

                for query in demo_queries:
                    print(f"\nSearching for: '{query}'")

                    # Method 1: Simple semantic search
                    print("  Simple Semantic Search:")
                    simple_results = summarizer.semantic_search_across_transcripts(
                        query, top_k=3
                    )
                    if simple_results:
                        for i, result in enumerate(
                            simple_results[:2]
                        ):  # Show top 2 results
                            print(
                                f"    Result {i+1} (Score: {result['similarity_score']:.3f}):"
                            )
                            print(
                                f"      Source: {result['metadata'].get('transcript_name', 'Unknown')}"
                            )
                            print(
                                f"      Type: {result['metadata'].get('chunk_type', 'Unknown')}"
                            )
                            print(f"      Text: {result['text'][:150]}...")
                    else:
                        print("    No relevant results found.")

                    # Method 2: Hierarchical retrieval with reranking
                    print("  Hierarchical Retrieval with Reranking:")
                    try:
                        hierarchical_results = (
                            summarizer.hierarchical_retrieval_with_reranking(
                                query,
                                top_k_per_transcript=3,
                                final_top_k=3,
                                use_llm_reranking=True,
                            )
                        )

                        if hierarchical_results:
                            for i, result in enumerate(hierarchical_results[:2]):
                                print(
                                    f"    Result {i+1} (Rank: {result.get('rank', 'N/A')}, "
                                    f"Final Score: {result.get('final_score', 'N/A'):.3f}):"
                                )
                                print(
                                    f"      Source: {result.get('transcript_name', 'Unknown')}"
                                )
                                print(
                                    f"      Type: {result.get('chunk_type', 'unknown')}"
                                )
                                print(f"      Text: {result.get('text', '')[:150]}...")
                        else:
                            print("    No relevant results found.")

                    except Exception as e:
                        print(f"    Error in hierarchical retrieval: {e}")

            print(f"\nAll files have been saved to the current directory.")
            print(f"Vector database stored at: {chroma_dir}")
        else:
            print("No transcripts were successfully processed.")

    except Exception as e:
        print(f"An error occurred: {e}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
