#!/usr/bin/env python3
"""
Document Processor for Transcript Analysis

This module handles reading and processing different document formats including
Word documents (.docx) and PDF files (.pdf) for transcript analysis.
"""

import os
from pathlib import Path
from typing import List, Optional, Dict, Any
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles reading and processing of different document formats."""

    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = [".docx", ".pdf"]
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if required dependencies are available."""
        # Check for python-docx (Word documents)
        try:
            import docx

            self.docx_available = True
            logger.info("python-docx available for Word document processing")
        except ImportError:
            self.docx_available = False
            logger.warning(
                "python-docx not available. Word documents cannot be processed."
            )

        # Check for PyPDF2 or pypdf (PDF files)
        try:
            import PyPDF2

            self.pypdf_available = True
            logger.info("PyPDF2 available for PDF processing")
        except ImportError:
            try:
                import pypdf

                self.pypdf_available = True
                logger.info("pypdf available for PDF processing")
            except ImportError:
                self.pypdf_available = False
                logger.warning(
                    "PyPDF2/pypdf not available. PDF files cannot be processed."
                )

    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        available_formats = []
        if self.docx_available:
            available_formats.append(".docx")
        if self.pypdf_available:
            available_formats.append(".pdf")
        return available_formats

    def can_process_file(self, file_path: str) -> bool:
        """Check if a file can be processed."""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats

    def extract_text_from_document(self, file_path: str) -> Optional[str]:
        """Extract text from a document file (Word or PDF)."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        file_ext = Path(file_path).suffix.lower()

        if file_ext == ".docx":
            return self._extract_text_from_docx(file_path)
        elif file_ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        else:
            logger.error(f"Unsupported file format: {file_ext}")
            return None

    def _extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from a Word document."""
        if not self.docx_available:
            logger.error("python-docx not available for Word document processing")
            return None

        try:
            from docx import Document

            doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            full_text = "\n".join(text_parts)
            logger.info(
                f"Successfully extracted {len(full_text)} characters from Word document: {file_path}"
            )
            return full_text

        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            return None

    def _extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from a PDF file."""
        if not self.pypdf_available:
            logger.error("PyPDF2/pypdf not available for PDF processing")
            return None

        try:
            # Try PyPDF2 first
            if hasattr(self, "pypdf_available") and self.pypdf_available:
                try:
                    import PyPDF2

                    with open(file_path, "rb") as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text_parts = []

                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text = page.extract_text()
                            if text.strip():
                                text_parts.append(text.strip())

                        full_text = "\n".join(text_parts)
                        logger.info(
                            f"Successfully extracted {len(full_text)} characters from PDF using PyPDF2: {file_path}"
                        )
                        return full_text

                except ImportError:
                    pass

                # Try pypdf
                try:
                    import pypdf

                    with open(file_path, "rb") as file:
                        pdf_reader = pypdf.PdfReader(file)
                        text_parts = []

                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text = page.extract_text()
                            if text.strip():
                                text_parts.append(text.strip())

                        full_text = "\n".join(text_parts)
                        logger.info(
                            f"Successfully extracted {len(full_text)} characters from PDF using pypdf: {file_path}"
                        )
                        return full_text

                except ImportError:
                    pass

            logger.error("No PDF processing library available")
            return None

        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return None

    def get_document_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get basic information about a document."""
        if not os.path.exists(file_path):
            return None

        try:
            file_stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()

            info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_stat.st_size,
                "file_extension": file_ext,
                "can_process": self.can_process_file(file_path),
                "supported_format": file_ext in self.supported_formats,
            }

            # Add format-specific info
            if file_ext == ".docx" and self.docx_available:
                try:
                    from docx import Document

                    doc = Document(file_path)
                    info["paragraphs"] = len(doc.paragraphs)
                    info["tables"] = len(doc.tables)
                except:
                    pass

            elif file_ext == ".pdf" and self.pypdf_available:
                try:
                    # Try to get page count
                    if hasattr(self, "pypdf_available") and self.pypdf_available:
                        try:
                            import PyPDF2

                            with open(file_path, "rb") as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                info["pages"] = len(pdf_reader.pages)
                        except:
                            try:
                                import pypdf

                                with open(file_path, "rb") as file:
                                    pdf_reader = pypdf.PdfReader(file)
                                    info["pages"] = len(pdf_reader.pages)
                            except:
                                pass
                except:
                    pass

            return info

        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {e}")
            return None

    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Process all supported documents in a directory."""
        if not os.path.exists(directory_path):
            logger.error(f"Directory not found: {directory_path}")
            return []

        results = []
        supported_formats = self.get_supported_formats()

        logger.info(f"Processing directory: {directory_path}")
        logger.info(f"Supported formats: {', '.join(supported_formats)}")

        for file_path in Path(directory_path).glob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_formats:
                logger.info(f"Processing file: {file_path.name}")

                # Get document info
                doc_info = self.get_document_info(str(file_path))
                if doc_info:
                    # Extract text
                    text = self.extract_text_from_document(str(file_path))
                    if text:
                        doc_info["text_length"] = len(text)
                        doc_info["text_preview"] = (
                            text[:200] + "..." if len(text) > 200 else text
                        )
                        doc_info["extraction_successful"] = True
                    else:
                        doc_info["extraction_successful"] = False
                        doc_info["text_length"] = 0
                        doc_info["text_preview"] = ""

                    results.append(doc_info)
                else:
                    logger.warning(f"Could not get info for file: {file_path.name}")

        logger.info(f"Processed {len(results)} documents from directory")
        return results

    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """Validate a document and provide processing status."""
        validation_result = {
            "file_path": file_path,
            "file_exists": False,
            "supported_format": False,
            "can_process": False,
            "dependencies_available": False,
            "validation_message": "",
        }

        if not os.path.exists(file_path):
            validation_result["validation_message"] = "File does not exist"
            return validation_result

        validation_result["file_exists"] = True

        file_ext = Path(file_path).suffix.lower()
        if file_ext in self.supported_formats:
            validation_result["supported_format"] = True
        else:
            validation_result["validation_message"] = f"Unsupported format: {file_ext}"
            return validation_result

        if file_ext == ".docx":
            validation_result["dependencies_available"] = self.docx_available
            if not self.docx_available:
                validation_result["validation_message"] = (
                    "python-docx library not available"
                )
                return validation_result
        elif file_ext == ".pdf":
            validation_result["dependencies_available"] = self.pypdf_available
            if not self.pypdf_available:
                validation_result["validation_message"] = (
                    "PDF processing libraries not available"
                )
                return validation_result

        validation_result["can_process"] = True
        validation_result["validation_message"] = (
            "Document can be processed successfully"
        )

        return validation_result


# Example usage and testing
if __name__ == "__main__":
    processor = DocumentProcessor()

    print("=== Document Processor Test ===")
    print(f"Supported formats: {processor.get_supported_formats()}")

    # Test with a sample directory
    test_dir = "FlexXray Transcripts"
    if os.path.exists(test_dir):
        print(f"\nProcessing directory: {test_dir}")
        results = processor.process_directory(test_dir)

        for result in results:
            print(f"\nFile: {result['file_name']}")
            print(f"  Format: {result['file_extension']}")
            print(f"  Size: {result['file_size']} bytes")
            print(f"  Can process: {result['can_process']}")
            if result.get("extraction_successful"):
                print(f"  Text length: {result['text_length']} characters")
                print(f"  Preview: {result['text_preview']}")
            else:
                print("  Text extraction failed")
    else:
        print(f"\nTest directory '{test_dir}' not found")
        print("Create the directory and add some documents to test")
